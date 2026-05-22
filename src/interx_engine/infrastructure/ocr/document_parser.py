"""
OCR & 문서 파싱 모듈 — 정부지원사업 첨부파일 텍스트 추출

Phase 1: pdfplumber (텍스트 기반 PDF) — Streamlit Cloud 호환
Phase 2: pytesseract + pdf2image (스캔 PDF OCR) — 로컬/Tesseract 필요
Phase 3: HWP/DOCX 파싱 — python-docx / olefile

사용법:
    from interx_engine.infrastructure.ocr.document_parser import extract_text_from_url
    text = extract_text_from_url("https://example.com/notice.pdf")
"""
from __future__ import annotations

import io
import logging
import re
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple

log = logging.getLogger("interx.ocr")

# ── 지원 확장자 ──────────────────────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".hwp", ".hwpx"}

# ── 추출 결과 ────────────────────────────────────────────────────────────────
class ExtractionResult:
    """문서 텍스트 추출 결과."""
    def __init__(self, text: str = "", method: str = "", pages: int = 0,
                 tables: Optional[List[List[List[str]]]] = None,
                 error: str = ""):
        self.text = text
        self.method = method     # "pdfplumber" / "pypdf" / "tesseract" / "docx" / "hwp"
        self.pages = pages
        self.tables = tables or []
        self.error = error

    @property
    def success(self) -> bool:
        return bool(self.text and len(self.text.strip()) > 20)

    def __repr__(self):
        return f"ExtractionResult(method={self.method}, pages={self.pages}, chars={len(self.text)}, ok={self.success})"


# ═══════════════════════════════════════════════════════════════════════════════
#  Phase 1: pdfplumber — 텍스트 기반 PDF (Streamlit Cloud 호환)
# ═══════════════════════════════════════════════════════════════════════════════

def extract_pdf_pdfplumber(file_bytes: bytes, max_pages: int = 20) -> ExtractionResult:
    """pdfplumber로 PDF 텍스트 + 테이블 추출."""
    try:
        import pdfplumber
    except ImportError:
        return ExtractionResult(error="pdfplumber 미설치")

    try:
        pdf = pdfplumber.open(io.BytesIO(file_bytes))
        texts = []
        tables_all = []
        page_count = min(len(pdf.pages), max_pages)

        for i in range(page_count):
            page = pdf.pages[i]
            # 텍스트 추출
            page_text = page.extract_text() or ""
            if page_text.strip():
                texts.append(page_text.strip())

            # 테이블 추출
            try:
                page_tables = page.extract_tables()
                if page_tables:
                    tables_all.extend(page_tables)
            except Exception:
                pass

        pdf.close()
        full_text = "\n\n".join(texts)

        if not full_text.strip():
            return ExtractionResult(
                text="", method="pdfplumber", pages=page_count,
                error="텍스트 없음 (스캔 PDF일 수 있음)"
            )

        return ExtractionResult(
            text=full_text[:8000],
            method="pdfplumber",
            pages=page_count,
            tables=tables_all[:10],
        )

    except Exception as e:
        return ExtractionResult(error=f"pdfplumber 파싱 실패: {e}")


def extract_pdf_pypdf(file_bytes: bytes, max_pages: int = 20) -> ExtractionResult:
    """pypdf fallback — pdfplumber 실패 시."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return ExtractionResult(error="pypdf 미설치")

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = []
        page_count = min(len(reader.pages), max_pages)

        for i in range(page_count):
            page_text = reader.pages[i].extract_text() or ""
            if page_text.strip():
                texts.append(page_text.strip())

        full_text = "\n\n".join(texts)
        return ExtractionResult(
            text=full_text[:8000],
            method="pypdf",
            pages=page_count,
        )
    except Exception as e:
        return ExtractionResult(error=f"pypdf 파싱 실패: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
#  Phase 2: pytesseract OCR — 스캔 PDF (로컬 전용)
# ═══════════════════════════════════════════════════════════════════════════════

def _check_tesseract() -> bool:
    """Tesseract OCR 엔진 설치 여부 확인."""
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def extract_pdf_ocr(file_bytes: bytes, max_pages: int = 5, lang: str = "kor+eng") -> ExtractionResult:
    """pytesseract + pdf2image로 스캔 PDF OCR."""
    if not _check_tesseract():
        return ExtractionResult(error="Tesseract OCR 미설치 (로컬 전용 기능)")

    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError as e:
        return ExtractionResult(error=f"OCR 패키지 미설치: {e}")

    try:
        # PDF → 이미지 변환
        images = convert_from_bytes(
            file_bytes,
            first_page=1,
            last_page=max_pages,
            dpi=200,
        )

        texts = []
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang=lang)
            if page_text.strip():
                texts.append(page_text.strip())
            log.debug("[OCR] 페이지 %d/%d 완료 (%d자)", i + 1, len(images), len(page_text))

        full_text = "\n\n".join(texts)
        return ExtractionResult(
            text=full_text[:8000],
            method="tesseract",
            pages=len(images),
        )
    except Exception as e:
        return ExtractionResult(error=f"OCR 실패: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
#  Phase 3: DOCX / HWP 파싱
# ═══════════════════════════════════════════════════════════════════════════════

def extract_docx(file_bytes: bytes) -> ExtractionResult:
    """python-docx로 .docx 텍스트 추출."""
    try:
        from docx import Document
    except ImportError:
        return ExtractionResult(error="python-docx 미설치")

    try:
        doc = Document(io.BytesIO(file_bytes))
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())

        # 테이블에서도 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)

        full_text = "\n".join(texts)
        return ExtractionResult(
            text=full_text[:8000],
            method="docx",
            pages=1,
        )
    except Exception as e:
        return ExtractionResult(error=f"DOCX 파싱 실패: {e}")


def extract_hwp(file_bytes: bytes) -> ExtractionResult:
    """
    HWP 파일 텍스트 추출.
    방법 1: olefile + zlib (바이너리 HWP 직접 파싱)
    방법 2: python-hwp 패키지 (설치 시)
    """
    # 방법 1: olefile 직접 파싱 (가장 호환성 좋음)
    try:
        import olefile
        import zlib

        ole = olefile.OleFileIO(io.BytesIO(file_bytes))

        # HWP 바이너리 형식: BodyText/Section0, Section1, ... 에 텍스트 저장
        texts = []
        section_idx = 0
        while True:
            stream_name = f"BodyText/Section{section_idx}"
            if not ole.exists(stream_name):
                break
            try:
                data = ole.openstream(stream_name).read()
                # HWP는 zlib 압축된 레코드 형식
                try:
                    decompressed = zlib.decompress(data, -15)
                except zlib.error:
                    decompressed = data

                # 바이너리에서 한글 텍스트 추출 (UTF-16LE)
                text_parts = _extract_text_from_hwp_binary(decompressed)
                if text_parts:
                    texts.append(text_parts)
            except Exception as e:
                log.debug("[HWP] Section%d 파싱 실패: %s", section_idx, e)
            section_idx += 1

        ole.close()

        if texts:
            full_text = "\n".join(texts)
            return ExtractionResult(
                text=full_text[:8000],
                method="hwp_olefile",
                pages=section_idx,
            )
    except ImportError:
        log.debug("[HWP] olefile 미설치")
    except Exception as e:
        log.debug("[HWP] olefile 파싱 실패: %s", e)

    return ExtractionResult(error="HWP 파싱 실패 (olefile 미설치 또는 비호환 형식)")


def _extract_text_from_hwp_binary(data: bytes) -> str:
    """
    HWP 바이너리 레코드에서 텍스트를 추출한다.
    HWP 형식: 4바이트 헤더(태그ID+레벨+사이즈) + 데이터
    텍스트 레코드 (HWPTAG_PARA_TEXT = 67): UTF-16LE 인코딩
    """
    texts = []
    i = 0
    while i + 4 <= len(data):
        # 레코드 헤더: 4바이트
        header = int.from_bytes(data[i:i+4], 'little')
        tag_id = header & 0x3FF
        # level = (header >> 10) & 0x3FF
        size = (header >> 20) & 0xFFF
        if size == 0xFFF:
            # 확장 사이즈: 다음 4바이트
            if i + 8 > len(data):
                break
            size = int.from_bytes(data[i+4:i+8], 'little')
            i += 8
        else:
            i += 4

        if i + size > len(data):
            break

        record_data = data[i:i+size]
        i += size

        # HWPTAG_PARA_TEXT = 67 (0x43)
        if tag_id == 67 and size >= 2:
            try:
                # UTF-16LE 디코딩, 제어 문자 필터링
                text = record_data.decode('utf-16-le', errors='ignore')
                # HWP 제어 문자 제거 (0x00~0x1F, 특수 HWP 마커)
                clean = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
                clean = clean.strip()
                if clean and len(clean) >= 2:
                    texts.append(clean)
            except Exception:
                pass

    return " ".join(texts)


# ═══════════════════════════════════════════════════════════════════════════════
#  통합 추출 함수
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text(file_bytes: bytes, filename: str, max_pages: int = 20) -> ExtractionResult:
    """
    파일 확장자에 따라 적절한 파서를 선택해 텍스트를 추출한다.

    추출 우선순위 (PDF):
      1. pdfplumber (텍스트 기반)
      2. pypdf (fallback)
      3. pytesseract OCR (스캔 PDF, Tesseract 설치 시)

    Args:
        file_bytes: 파일 바이너리 데이터
        filename: 파일명 (확장자 판별용)
        max_pages: 최대 처리 페이지 수
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        # Phase 1: pdfplumber
        result = extract_pdf_pdfplumber(file_bytes, max_pages)
        if result.success:
            log.info("[OCR] PDF 텍스트 추출 성공: %s (%s, %d자)",
                     filename, result.method, len(result.text))
            return result

        # Phase 1 fallback: pypdf
        result2 = extract_pdf_pypdf(file_bytes, max_pages)
        if result2.success:
            log.info("[OCR] PDF pypdf fallback 성공: %s (%d자)", filename, len(result2.text))
            return result2

        # Phase 2: OCR (Tesseract 설치 시만)
        if _check_tesseract():
            result3 = extract_pdf_ocr(file_bytes, min(max_pages, 5))
            if result3.success:
                log.info("[OCR] PDF OCR 성공: %s (%d자)", filename, len(result3.text))
                return result3

        # 모든 방법 실패
        return ExtractionResult(
            error=f"PDF 텍스트 추출 실패: {result.error or result2.error}",
            method="none",
        )

    elif ext == ".docx":
        result = extract_docx(file_bytes)
        if result.success:
            log.info("[OCR] DOCX 추출 성공: %s (%d자)", filename, len(result.text))
        return result

    elif ext in (".hwp", ".hwpx"):
        result = extract_hwp(file_bytes)
        if result.success:
            log.info("[OCR] HWP 추출 성공: %s (%d자)", filename, len(result.text))
        return result

    elif ext == ".doc":
        # .doc는 바이너리 형식 → olefile 시도
        return ExtractionResult(error=".doc 형식 미지원 (DOCX로 변환 필요)")

    else:
        return ExtractionResult(error=f"지원하지 않는 확장자: {ext}")


def extract_text_from_url(
    url: str,
    filename: str = "",
    timeout: int = 15,
    max_size_mb: int = 10,
) -> ExtractionResult:
    """
    URL에서 파일을 다운로드한 후 텍스트를 추출한다.

    Args:
        url: 파일 다운로드 URL
        filename: 파일명 (없으면 URL에서 추출)
        timeout: 다운로드 타임아웃 (초)
        max_size_mb: 최대 파일 크기 (MB)
    """
    import requests

    if not filename:
        # URL에서 파일명 추출
        from urllib.parse import urlparse, unquote
        path = urlparse(url).path
        filename = unquote(path.split("/")[-1]) if path else "unknown"

    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return ExtractionResult(error=f"지원하지 않는 확장자: {ext}")

    try:
        resp = requests.get(
            url,
            timeout=timeout,
            headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                              "AppleWebKit/537.36 Chrome/124.0 Safari/537.36"
            },
            stream=True,
            verify=False,
        )
        resp.raise_for_status()

        # Content-Length 체크
        content_length = int(resp.headers.get("content-length", 0))
        if content_length > max_size_mb * 1024 * 1024:
            return ExtractionResult(error=f"파일 크기 초과: {content_length / 1024 / 1024:.1f}MB > {max_size_mb}MB")

        file_bytes = resp.content
        if len(file_bytes) > max_size_mb * 1024 * 1024:
            return ExtractionResult(error=f"파일 크기 초과: {len(file_bytes) / 1024 / 1024:.1f}MB")

        return extract_text(file_bytes, filename)

    except requests.exceptions.Timeout:
        return ExtractionResult(error=f"다운로드 타임아웃: {timeout}초")
    except requests.exceptions.RequestException as e:
        return ExtractionResult(error=f"다운로드 실패: {e}")
    except Exception as e:
        return ExtractionResult(error=f"파일 처리 실패: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
#  배치 추출 (공고의 attachment_items에서 일괄 추출)
# ═══════════════════════════════════════════════════════════════════════════════

def extract_from_attachments(
    attachment_items: List[Dict[str, str]],
    max_files: int = 3,
    max_pages: int = 10,
) -> Tuple[str, List[Dict[str, str]]]:
    """
    공고의 첨부파일 목록에서 텍스트 추출.

    Args:
        attachment_items: [{"name": "공고문.pdf", "url": "https://..."}, ...]
        max_files: 최대 처리 파일 수
        max_pages: PDF 최대 페이지 수

    Returns:
        (combined_text, extraction_logs)
        combined_text: 추출된 텍스트 합산 (body_text 보강용)
        extraction_logs: [{"name": ..., "method": ..., "chars": ..., "error": ...}, ...]
    """
    texts = []
    logs = []
    processed = 0

    for item in attachment_items:
        if processed >= max_files:
            break

        name = item.get("name", "")
        url = item.get("url", "")

        if not url:
            continue

        ext = Path(name).suffix.lower() if name else ""
        # URL에서도 확장자 추출 시도
        if not ext:
            from urllib.parse import urlparse
            ext = Path(urlparse(url).path).suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            continue

        log.info("[OCR] 첨부파일 처리: %s (%s)", name[:50], ext)
        result = extract_text_from_url(url, name or f"attachment{ext}")
        processed += 1

        log_entry = {
            "name": name[:100],
            "url": url[:200],
            "method": result.method,
            "chars": len(result.text),
            "pages": result.pages,
            "success": result.success,
            "error": result.error,
        }
        logs.append(log_entry)

        if result.success:
            texts.append(f"[첨부: {name[:50]}]\n{result.text}")
            log.info("[OCR] 추출 성공: %s → %d자 (%s)", name[:40], len(result.text), result.method)
        else:
            log.debug("[OCR] 추출 실패: %s → %s", name[:40], result.error)

    combined = "\n\n".join(texts)
    return combined[:8000], logs
