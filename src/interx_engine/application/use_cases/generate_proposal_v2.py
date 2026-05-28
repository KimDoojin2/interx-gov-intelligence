"""
제안서 초안 자동 생성 v2 — 솔루션 맞춤 + 경쟁 분석 + 정기공고 이력 반영
python-docx 기반, A/B 등급 공고 대상
"""
from __future__ import annotations

import logging
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional

from interx_engine.core.entities.notice import Notice
from interx_engine.core.entities.score_card import ScoreCard

log = logging.getLogger("interx.proposal_v2")

# ── 솔루션별 상세 역량 설명 ──────────────────────────────────────────────────
SOLUTION_PROFILES = {
    "ManufacturingDT": {
        "name": "제조 디지털트윈",
        "desc": "실시간 공정 시뮬레이션 및 최적화 플랫폼",
        "strengths": [
            "3D 공정 시뮬레이션 기반 디지털트윈 플랫폼 구축 경험",
            "실시간 센서 데이터 수집 → 가상 공장 동기화 기술",
            "공정 변경 시나리오 사전 시뮬레이션으로 리스크 최소화",
        ],
        "tech_stack": "Unity 3D / Unreal Engine / OPC-UA / MQTT / Azure IoT",
    },
    "RecipeAI": {
        "name": "레시피 AI",
        "desc": "공정 조건 최적화 및 배합 예측 AI",
        "strengths": [
            "다변량 공정 파라미터 최적화 AI 모델 개발 경험",
            "레시피 자동 추천 및 품질 예측 정확도 95%+",
            "실시간 공정 조건 피드백 루프 구현",
        ],
        "tech_stack": "XGBoost / LSTM / Bayesian Optimization / MLOps",
    },
    "QualityAI": {
        "name": "품질 AI",
        "desc": "불량 검출 및 품질 예측 시스템",
        "strengths": [
            "제조 불량 자동 검출 AI (정확도 99%+)",
            "SPC 통계적 공정 관리 자동화",
            "품질 이상 조기 경보 시스템 구축",
        ],
        "tech_stack": "CNN / YOLO / Anomaly Detection / Edge AI",
    },
    "InspectionAI": {
        "name": "비전 검사 AI",
        "desc": "머신비전 기반 자동 외관 검사",
        "strengths": [
            "멀티 카메라 비전 검사 시스템 구축 경험 다수",
            "미세 결함(0.1mm 이하) 탐지 딥러닝 모델",
            "라인 속도 대응 실시간 검사 (100ms 이내 판정)",
        ],
        "tech_stack": "YOLO v8 / Detectron2 / OpenCV / GigE Vision",
    },
    "SafetyAI": {
        "name": "안전 AI",
        "desc": "중대재해 예방 및 작업자 안전 모니터링",
        "strengths": [
            "중대재해처벌법 대응 AI 안전 관리 솔루션",
            "CCTV 기반 위험 행동 실시간 탐지",
            "IoT 센서 연계 위험 환경 자동 경보",
        ],
        "tech_stack": "Pose Estimation / Object Detection / IoT Gateway",
    },
    "GenAI": {
        "name": "제조 GenAI",
        "desc": "생성형 AI 기반 제조 지식 자동화",
        "strengths": [
            "제조 도메인 특화 LLM/RAG 시스템 구축",
            "설비 매뉴얼 자동 QA 챗봇 개발 경험",
            "공정 데이터 기반 자동 보고서 생성",
        ],
        "tech_stack": "GPT-4 / Claude / LangChain / RAG / Vector DB",
    },
    "InfraDS": {
        "name": "데이터 인프라",
        "desc": "제조 데이터 스페이스 및 클라우드 인프라",
        "strengths": [
            "Catena-X / AAS 기반 데이터 스페이스 구축 경험",
            "제조 데이터 레이크 설계 및 ETL 파이프라인",
            "클라우드-엣지 하이브리드 아키텍처",
        ],
        "tech_stack": "Catena-X / Eclipse Dataspace / K8s / Kafka / MinIO",
    },
    "PdM": {
        "name": "예지보전",
        "desc": "설비 고장 예측 및 최적 정비 계획",
        "strengths": [
            "진동/온도/전류 데이터 기반 고장 예측 모델",
            "설비 잔여 수명(RUL) 예측 정확도 90%+",
            "정비 비용 30%+ 절감 실적",
        ],
        "tech_stack": "LSTM / Transformer / PHM / Edge Computing",
    },
}


def _calc_dday(deadline: str) -> int:
    try:
        from datetime import datetime
        dl = datetime.strptime(deadline, "%Y-%m-%d").date()
        return (dl - date.today()).days
    except (ValueError, TypeError):
        return -1


def generate_proposals_v2(
    notices: List[Notice],
    score_cards: List[ScoreCard],
    output_dir: str = "",
    target_grades: tuple = ("A", "B"),
) -> List[str]:
    """
    A/B 등급 공고 제안서 초안 v2 — 솔루션 맞춤 + 상세 역량.
    Returns: 생성된 파일 경로 목록
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.warning("[ProposalV2] python-docx 미설치 — pip install python-docx")
        return []

    score_map = {s.notice_id: s for s in score_cards}

    if not output_dir:
        try:
            from interx_engine.application.ports.settings_port import proposal_output_dir
            output_dir = proposal_output_dir()
        except Exception:
            import tempfile
            output_dir = str(Path(tempfile.gettempdir()) / "interx_proposals")

    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    generated = []

    for notice in notices:
        sc = score_map.get(notice.notice_id)
        if not sc or sc.priority_grade not in target_grades:
            continue

        doc = Document()

        # ── 헤더 ─────────────────────────────────────────────────────────────
        h = doc.add_heading(f"[제안서 초안] {notice.title}", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x00, 0xCF, 0xFF)

        # ── 핵심 지표 카드 ───────────────────────────────────────────────────
        dday = _calc_dday(notice.deadline_date or "")
        badge_text = (
            f"  {sc.priority_grade} 등급  |  "
            f"적합도: {sc.fitness_score:.0f}점  |  "
            f"우선순위: {sc.priority_score:.0f}점  |  "
            f"D-day: {dday}일"
        )
        p = doc.add_paragraph()
        run = p.add_run(badge_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E) if sc.priority_grade == "A" else RGBColor(0x00, 0xCF, 0xFF)

        # L3 / 정기공고 플래그
        flags = []
        if getattr(notice, "l3_strong", "N") == "Y":
            flags.append("L3 Strong")
        if getattr(notice, "recurring_flag", "N") == "Y":
            flags.append(f"Recurring: {getattr(notice, 'recurring_group', '')}")
        if flags:
            fp = doc.add_paragraph()
            fr = fp.add_run("  " + " | ".join(flags))
            fr.font.color.rgb = RGBColor(0xFF, 0xD7, 0x00)
            fr.bold = True

        doc.add_paragraph()

        # ── 1. 사업 개요 ─────────────────────────────────────────────────────
        doc.add_heading("1. 사업 개요", level=2)
        overview = (
            f"공고명:     {notice.title}\n"
            f"주관기관:   {notice.agency or '-'}\n"
            f"주무부처:   {notice.ministry or '-'}\n"
            f"공고일:     {notice.posted_date or '-'}\n"
            f"마감일:     {notice.deadline_date or '-'} (D-{dday})\n"
            f"예산:       {notice.budget or '-'}\n"
            f"사이트:     {notice.site}"
        )
        doc.add_paragraph(overview)

        # ── 2. 적합도 분석 ───────────────────────────────────────────────────
        doc.add_heading("2. InterX 적합도 분석", level=2)
        keywords_str = " / ".join(sc.positive_keywords[:8]) if sc.positive_keywords else "-"
        analysis = (
            f"적합도 점수:   {sc.fitness_score:.1f} / 100\n"
            f"우선순위 점수: {sc.priority_score:.1f} / 100\n"
            f"산업 적합도:   {sc.industry_score:.1f} / 100\n"
            f"매칭 키워드:   {keywords_str}\n"
        )
        doc.add_paragraph(analysis)

        # 솔루션 점수 테이블
        if sc.solution_scores:
            active = {k: v for k, v in sc.solution_scores.items() if v > 0}
            if active:
                doc.add_paragraph("솔루션별 적합 점수:")
                table = doc.add_table(rows=1, cols=3)
                table.style = "Light List Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text = "솔루션"
                hdr[1].text = "점수"
                hdr[2].text = "적합도"
                for sol_name, sol_score in sorted(active.items(), key=lambda x: -x[1]):
                    row = table.add_row().cells
                    profile = SOLUTION_PROFILES.get(sol_name, {})
                    row[0].text = profile.get("name", sol_name)
                    row[1].text = f"{sol_score:.0f}"
                    row[2].text = "HIGH" if sol_score >= 50 else "MID" if sol_score >= 25 else "LOW"

        doc.add_paragraph()

        # ── 3. 추천 솔루션 상세 ──────────────────────────────────────────────
        doc.add_heading("3. InterX 추천 솔루션 및 역량", level=2)

        recommended = []
        if sc.solution_scores:
            recommended = sorted(
                [(k, v) for k, v in sc.solution_scores.items() if v > 0],
                key=lambda x: -x[1],
            )[:3]  # 상위 3개 솔루션

        if recommended:
            for sol_key, sol_score in recommended:
                profile = SOLUTION_PROFILES.get(sol_key, {})
                if profile:
                    doc.add_heading(
                        f"3-{recommended.index((sol_key, sol_score))+1}. "
                        f"{profile['name']} ({sol_key})",
                        level=3,
                    )
                    doc.add_paragraph(f"설명: {profile['desc']}")
                    doc.add_paragraph("핵심 역량:")
                    for s in profile["strengths"]:
                        doc.add_paragraph(f"  - {s}")
                    doc.add_paragraph(f"기술 스택: {profile['tech_stack']}")
                    doc.add_paragraph()
        else:
            doc.add_paragraph("[솔루션 점수 데이터 없음 — 수동 작성 필요]")

        # ── 4. 제안 전략 ─────────────────────────────────────────────────────
        doc.add_heading("4. 제안 전략", level=2)
        action = getattr(notice, "recommended_action", "") or "제안 검토"
        strategy = (
            f"추천 액션: {action}\n\n"
            f"[작성 필요] 본 사업의 차별화 포인트:\n"
            f"  1. 기술적 차별화: ___\n"
            f"  2. 가격 경쟁력: ___\n"
            f"  3. 수행 실적: ___\n"
            f"  4. 컨소시엄 구성: ___"
        )
        doc.add_paragraph(strategy)

        # ── 5. 경쟁 분석 (있는 경우) ─────────────────────────────────────────
        comp_flag = getattr(notice, "competitor_flag", "")
        if comp_flag:
            doc.add_heading("5. 경쟁 환경 분석", level=2)
            doc.add_paragraph(f"감지된 경쟁사: {comp_flag}")
            doc.add_paragraph(
                "[작성 필요] 경쟁사 대비 InterX 강점 분석:\n"
                "  - 기술 차별점: ___\n"
                "  - 가격 차별점: ___\n"
                "  - 실적 차별점: ___"
            )

        # ── 6. 추진 일정 ─────────────────────────────────────────────────────
        doc.add_heading("6. 추진 일정", level=2)
        timeline = (
            f"제안서 제출 기한: {notice.deadline_date or '-'}\n"
            f"담당자: {notice.manager or '미배정'}\n"
            f"BD 마일스톤: {notice.bd_milestone or '-'}\n\n"
            f"[작성 필요]\n"
            f"  - 제안서 작성: ~ D-7\n"
            f"  - 내부 리뷰: D-5\n"
            f"  - 최종 수정: D-3\n"
            f"  - 제출: D-1"
        )
        doc.add_paragraph(timeline)

        # ── 7. 공고 원문 ─────────────────────────────────────────────────────
        doc.add_heading("7. 공고 원문 링크", level=2)
        url = notice.detail_url or getattr(notice, "notice_link", "") or "-"
        doc.add_paragraph(url)

        # ── 8. 공고 요약 (있는 경우) ─────────────────────────────────────────
        summary = getattr(notice, "summary", "") or ""
        body = getattr(notice, "body_text", "") or ""
        if summary or body:
            doc.add_heading("8. 공고 본문 요약", level=2)
            if summary:
                doc.add_paragraph(summary)
            elif body:
                doc.add_paragraph(body[:2000] + ("..." if len(body) > 2000 else ""))

        # ── 저장 ─────────────────────────────────────────────────────────────
        safe = "".join(c for c in notice.title[:30] if c.isalnum() or c in " _-")
        safe = safe.strip() or notice.notice_id
        fname = out_dir / f"{sc.priority_grade}_{notice.site}_{safe}.docx"
        doc.save(str(fname))
        generated.append(str(fname))
        log.info("[ProposalV2] %s", fname.name)

    log.info("[ProposalV2] %d개 제안서 생성 → %s", len(generated), out_dir)
    return generated
