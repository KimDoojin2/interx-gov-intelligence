"""
제안서 초안 자동 생성 — python-docx 기반
P1/P2 등급 공고에 대해 Word 파일 자동 생성
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import List

from interx_engine.core.entities.notice import Notice
from interx_engine.core.entities.score_card import ScoreCard

log = logging.getLogger("interx.proposal")

_SECTIONS = [
    ("1. 사업 개요",
     "공고명:   {title}\n주관기관: {agency}\n주무부처: {ministry}\n공고일:   {posted}\n마감일:   {deadline}"),
    ("2. 지원 규모",
     "예산: {budget}\n지원 기간: {duration}"),
    ("3. InterX 추천 솔루션",
     "{solution}\n\n핵심 적합 키워드: {keywords}"),
    ("4. InterX 강점",
     "• 제조 AI 전문 역량 (AX · 스마트팩토리 · 디지털트윈)\n"
     "• {keywords} 핵심 기술 보유\n"
     "• 유사 프로젝트 다수 수행 경험\n"
     "• [작성 필요] 구체적 수행 실적 기재"),
    ("5. 제안 전략",
     "추천 액션: {action}\n\n[작성 필요] 차별화 포인트 및 제안 핵심 내용을 여기에 기재하세요."),
    ("6. 추진 일정",
     "제안서 제출 기한: {deadline}\n담당자: {manager}\n\n[작성 필요] 세부 일정표 첨부"),
    ("7. 공고 원문 링크",
     "{url}"),
]


def _default_output_dir() -> str:
    try:
        from interx_engine.application.ports.settings_port import proposal_output_dir
        return proposal_output_dir()
    except Exception:
        import tempfile
        return str(Path(tempfile.gettempdir()) / "interx_proposals")


def generate_proposals(
    notices: List[Notice],
    score_cards: List[ScoreCard],
    output_dir: str = "",
    target_grades: tuple = ("A", "B"),
) -> List[str]:
    """
    A/B 등급 공고 제안서 초안 Word 파일 생성.
    Returns: 생성된 파일 경로 목록
    """
    try:
        from docx import Document                              # type: ignore
        from docx.shared import Pt, RGBColor                  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH         # type: ignore
    except ImportError:
        log.warning("[Proposal] python-docx 미설치 — !pip install python-docx")
        return []

    score_map = {s.notice_id: s for s in score_cards}
    out_dir   = Path(output_dir or _default_output_dir())
    out_dir.mkdir(parents=True, exist_ok=True)
    generated = []

    for notice in notices:
        score = score_map.get(notice.notice_id)
        if not score or score.priority_grade not in target_grades:
            continue

        doc = Document()

        # ── 제목 ─────────────────────────────────────────────────────────────
        h = doc.add_heading(f"[제안서 초안] {notice.title}", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x46, 0x8A)

        # ── 등급 뱃지 ─────────────────────────────────────────────────────────
        p    = doc.add_paragraph()
        run  = p.add_run(
            f"  ★ {score.priority_grade} 등급  |  "
            f"적합도: {score.fitness_score:.0f}점  |  "
            f"추천솔루션: {notice.recommended_solution}"
        )
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph()

        # ── 섹션 ─────────────────────────────────────────────────────────────
        for section_title, template in _SECTIONS:
            doc.add_heading(section_title, level=2)
            content = template.format(
                title    = notice.title,
                agency   = notice.agency        or "-",
                ministry = notice.ministry      or "-",
                posted   = notice.posted_date   or "-",
                deadline = notice.deadline_date or "-",
                budget   = notice.budget        or "-",
                duration = notice.duration_months or "-",
                solution = notice.recommended_solution or "-",
                keywords = " / ".join(score.positive_keywords[:5]) if score else "-",
                action   = notice.recommended_action or "제안 검토",
                manager  = notice.manager       or "미배정",
                url      = notice.detail_url    or notice.notice_link or "-",
            )
            doc.add_paragraph(content)

        # ── 저장 ─────────────────────────────────────────────────────────────
        safe = "".join(c for c in notice.title[:30] if c.isalnum() or c in " _-한글")
        fname = out_dir / f"{score.priority_grade}_{notice.site}_{safe}.docx"
        doc.save(str(fname))
        generated.append(str(fname))
        log.info("[Proposal] 생성: %s", fname.name)

    log.info("[Proposal] 총 %d개 제안서 초안 → %s", len(generated), out_dir)
    return generated
