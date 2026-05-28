"""
사업계획서 AI 자동 생성기 — 공고 맞춤형 DOCX 출력 (하이브리드) v3.

두 가지 모드 지원:
  MODE 1 (양식 업로드): HWP/HWPX/PDF 양식 → 섹션 추출 → AI가 내용 채움
  MODE 2 (공고 분석):   공고 본문에서 요구사항/목차 → AI가 구조+내용 생성

v3 핵심 개선:
  - configs/company_knowledge.yaml 기반 회사 지식베이스 활용
  - 섹션별 특화 프롬프트 (개요/필요성/목표/추진체계/사업화 등)
  - 시장 데이터, 경쟁사 정보, 기술 표준 자동 주입
  - max_tokens 4096으로 상세 콘텐츠 생성
  - fallback도 지식베이스 기반 풍부한 콘텐츠

핵심 원칙:
  - 사업마다 성격이 다르므로, 고정 템플릿 대신 공고별 동적 구조 생성
  - 비용/예산/금액 정보는 절대 포함하지 않음
  - Gemini 2.0 Flash 무료 API (15 RPM)
"""
from __future__ import annotations

import logging
import re
import yaml
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from interx_engine.core.entities.notice import Notice
from interx_engine.core.entities.score_card import ScoreCard

log = logging.getLogger("interx.business_plan")


# ═════════════════════════════════════════════════════════════════════════════
#  회사 지식베이스 로딩
# ═════════════════════════════════════════════════════════════════════════════

_KNOWLEDGE_BASE: Optional[Dict] = None

def _load_knowledge_base() -> Dict:
    """configs/company_knowledge.yaml 로드 (싱글턴)."""
    global _KNOWLEDGE_BASE
    if _KNOWLEDGE_BASE is not None:
        return _KNOWLEDGE_BASE

    # 프로젝트 루트에서 configs/ 찾기
    candidates = [
        Path(__file__).resolve().parents[4] / "configs" / "company_knowledge.yaml",
        Path(__file__).resolve().parents[3] / "configs" / "company_knowledge.yaml",
        Path(__file__).resolve().parents[2] / "configs" / "company_knowledge.yaml",
    ]
    for p in candidates:
        if p.exists():
            try:
                with open(p, "r", encoding="utf-8") as f:
                    _KNOWLEDGE_BASE = yaml.safe_load(f) or {}
                log.info("[BusinessPlan] 지식베이스 로드: %s", p.name)
                return _KNOWLEDGE_BASE
            except Exception as e:
                log.warning("[BusinessPlan] 지식베이스 로드 실패: %s", e)

    _KNOWLEDGE_BASE = {}
    return _KNOWLEDGE_BASE


def _get_company_context() -> str:
    """회사 기본 정보를 프롬프트 텍스트로 변환."""
    kb = _load_knowledge_base()
    company = kb.get("company", {})
    if not company:
        return "(주)인터엑스 — 제조 AI 전문기업"

    lines = [
        f"기업명: {company.get('name', '(주)인터엑스')}",
        f"설립: {company.get('established', '')}년 | 대표: {company.get('ceo', '')}",
        f"유형: {company.get('type', '')}",
        f"인력: {company.get('employees', '')}명",
        "",
        "■ 핵심 역량:",
    ]
    for cap in company.get("core_competency", []):
        lines.append(f"  - {cap}")

    return "\n".join(lines)


def _get_solutions_context(solutions: List[str]) -> str:
    """관련 솔루션의 상세 기술 역량을 프롬프트 텍스트로 변환."""
    kb = _load_knowledge_base()
    sol_data = kb.get("solutions", {})
    lines = []

    # INTERX_CAPABILITIES 기본 정보
    for s in solutions:
        cap = INTERX_CAPABILITIES.get(s, {})
        if cap:
            lines.append(f"■ {cap['name']} ({s})")
            lines.append(f"  설명: {cap['desc']}")
            lines.append(f"  기술스택: {cap['tech']}")
            for st in cap.get("strengths", []):
                lines.append(f"  - {st}")
            lines.append("")

    # 지식베이스의 상세 솔루션 정보 추가
    sol_mapping = {
        "GenAI": ["multi_agent_platform", "ai_copilot", "ontology_knowledge"],
        "PdM": ["pdm_system"],
        "QualityAI": ["quality_ai"],
        "InspectionAI": ["quality_ai"],
        "RecipeAI": ["recipe_optimization"],
        "ManufacturingDT": ["digital_twin"],
        "InfraDS": ["ontology_knowledge"],
        "SafetyAI": [],
    }

    added = set()
    for s in solutions:
        for sol_key in sol_mapping.get(s, []):
            if sol_key in added:
                continue
            added.add(sol_key)
            detail = sol_data.get(sol_key, {})
            if detail:
                lines.append(f"■ 상세: {detail.get('name', sol_key)}")
                desc = detail.get("description", "")
                if desc:
                    lines.append(f"  {desc.strip()[:500]}")
                for feat in detail.get("key_features", [])[:5]:
                    lines.append(f"  - {feat}")

                # Agent 정보 (멀티-Agent 플랫폼)
                for agent in detail.get("agents", [])[:4]:
                    lines.append(f"  [{agent['name']}] {agent['role']}")
                    lines.append(f"    기술: {agent['tech']}")

                arch = detail.get("architecture", "")
                if arch:
                    lines.append(f"  아키텍처: {arch.strip()[:300]}")
                lines.append("")

    return "\n".join(lines) if lines else "범용 제조 AI 솔루션"


def _get_market_context() -> str:
    """시장 데이터를 프롬프트 텍스트로 변환."""
    kb = _load_knowledge_base()
    market = kb.get("market_data", {})
    if not market:
        return ""

    lines = ["■ 시장 데이터:"]

    dt = market.get("digital_twin", {})
    if dt:
        kr = dt.get("korea", {})
        gl = dt.get("global", {})
        if kr:
            lines.append(f"  [한국 DT 시장] {kr.get('size_2024','')} (2024) → {kr.get('size_2033','')} (2033), CAGR {kr.get('cagr','')}")
            lines.append(f"    출처: {kr.get('source','')}")
        if gl:
            lines.append(f"  [글로벌 DT 시장] {gl.get('size_2025','')} (2025) → {gl.get('size_2034','')} (2034), CAGR {gl.get('cagr','')}")

    ai = market.get("ai_agent", {})
    if ai:
        g = ai.get("global", {})
        if g:
            lines.append(f"  [글로벌 AI Agent 시장] {g.get('size_2024','')} (2024) → {g.get('size_2030','')} (2030), CAGR {g.get('cagr','')}")

    kai = market.get("korea_ai_industry", {})
    if kai:
        lines.append(f"  [한국 AI 산업] 총매출 {kai.get('total_revenue_2024','')}, B2B 비중 {kai.get('b2b_share','')}")
        lines.append(f"    에이전트 기술 비중: {kai.get('agent_tech_ratio','')}")

    return "\n".join(lines)


def _get_competitor_context() -> str:
    """경쟁사 정보를 프롬프트 텍스트로 변환."""
    kb = _load_knowledge_base()
    comp = kb.get("competitors", {})
    if not comp:
        return ""

    lines = ["■ 국내외 경쟁사 현황:"]
    for c in comp.get("domestic", [])[:4]:
        lines.append(f"  [{c['name']}] {c['product']} — {c['focus']}")
    for c in comp.get("global", [])[:3]:
        lines.append(f"  [{c['name']}] {c['product']} — {c['focus']}")

    diff = comp.get("interx_differentiation", "")
    if diff:
        lines.append(f"\n■ 인터엑스 차별점: {diff.strip()[:500]}")

    return "\n".join(lines)


def _get_standards_context() -> str:
    """기술 표준 정보를 프롬프트 텍스트로 변환."""
    kb = _load_knowledge_base()
    std = kb.get("standards", {})
    if not std:
        return ""

    lines = ["■ 관련 기술 표준:"]
    aas = std.get("aas", {})
    if aas:
        lines.append(f"  [AAS] {aas.get('standard','')} / 국내: {aas.get('korea_standard','')}")
        lines.append(f"    {aas.get('description','').strip()[:200]}")
        lines.append(f"    인터엑스 역량: {aas.get('interx_capability','').strip()[:200]}")

    for s in std.get("digital_twin_standards", []):
        lines.append(f"  - {s}")

    return "\n".join(lines)


def _get_track_record_context() -> str:
    """수행실적 정보를 프롬프트 텍스트로 변환."""
    kb = _load_knowledge_base()
    tr = kb.get("track_record", {})
    if not tr:
        return ""

    lines = ["■ 주요 수행실적:"]
    for p in tr.get("projects", []):
        lines.append(f"  [{p['name']}] {p.get('description','')}")
        if p.get("client"):
            lines.append(f"    고객: {p['client']}")
        for feat in p.get("features", [])[:3]:
            lines.append(f"    - {feat}")

    return "\n".join(lines)


# ═════════════════════════════════════════════════════════════════════════════
#  InterX 솔루션 역량 DB (기본)
# ═════════════════════════════════════════════════════════════════════════════

INTERX_CAPABILITIES = {
    "ManufacturingDT": {
        "name": "제조 디지털트윈",
        "desc": "실시간 공정 시뮬레이션 및 최적화 플랫폼",
        "keywords": ["디지털트윈", "시뮬레이션", "가상공장", "DT", "Digital Twin"],
        "tech": "Unity 3D / OPC-UA / MQTT / Azure IoT / 2D 블록 토폴로지",
        "strengths": [
            "2D 블록 기반 공정 토폴로지 구성 및 비주얼 트윈 플랫폼 구축 경험",
            "AI 에이전트 연동 What-if 시뮬레이션 기반 생산계획 사전 검증",
            "병목 및 큐타임 탐지(Bottleneck/Q-Time Detect) 알고리즘 구현",
            "레거시 시스템(MES/ERP/PLC) 파일 기반 에뮬레이션 연계 기술",
        ],
    },
    "RecipeAI": {
        "name": "레시피 AI / 공정조건 최적화",
        "desc": "공정 조건 최적화 및 배합 예측 AI, RMS 연동 자동 Recipe 보정",
        "keywords": ["레시피", "공정최적화", "배합", "파라미터", "Recipe"],
        "tech": "XGBoost / LSTM / Bayesian Optimization / MLOps / RMS 연동",
        "strengths": [
            "다변량 공정 파라미터 최적화 AI 모델 개발 경험",
            "레시피 자동 추천 및 품질 예측 정확도 95%+",
            "RMS 시스템 연동 Recipe 자동 반영 기술",
            "품질예측 결과 기반 실시간 절삭 조건 자동 조정",
        ],
    },
    "QualityAI": {
        "name": "품질 AI",
        "desc": "센서+공정데이터 결합 멀티모달 품질 예측, SPC 자동화",
        "keywords": ["불량검출", "품질예측", "SPC", "품질관리", "품질"],
        "tech": "CNN / Anomaly Detection / SPC / Feature Importance / Edge AI",
        "strengths": [
            "센서신호 + CNC운용데이터 + 공정조건 결합 멀티모달 품질 예측 모델",
            "온톨로지/FMEA 기반 Top-N 원인계 추정 기술",
            "SPC 통계적 공정 관리 자동화",
            "품질 불량 발생 확률(%) 실시간 추론",
        ],
    },
    "InspectionAI": {
        "name": "비전 검사 AI",
        "desc": "머신비전 기반 자동 외관 검사 (미세결함 0.1mm 탐지)",
        "keywords": ["비전검사", "외관검사", "머신비전", "비전", "검사"],
        "tech": "YOLO v8 / Detectron2 / OpenCV / GigE Vision",
        "strengths": [
            "멀티 카메라 비전 검사 시스템 구축 경험 다수",
            "미세 결함(0.1mm 이하) 탐지 딥러닝 모델",
            "라인 속도 대응 실시간 검사 (100ms 이내 판정)",
        ],
    },
    "SafetyAI": {
        "name": "안전 AI",
        "desc": "중대재해 예방 및 작업자 안전 모니터링",
        "keywords": ["안전", "중대재해", "재해예방", "안전관리"],
        "tech": "Pose Estimation / Object Detection / IoT Gateway",
        "strengths": [
            "중대재해처벌법 대응 AI 안전 관리 솔루션",
            "CCTV 기반 위험 행동 실시간 탐지",
            "IoT 센서 연계 위험 환경 자동 경보",
        ],
    },
    "GenAI": {
        "name": "제조 GenAI / 멀티-Agent 플랫폼",
        "desc": "산업 특화 멀티-Agent 오케스트레이션, Agentic AI Copilot",
        "keywords": ["생성형AI", "LLM", "챗봇", "RAG", "자연어", "Agent", "에이전트", "Copilot"],
        "tech": "LLM / LangChain / RAG / Vector DB / ReAct / Tool Calling / HITL",
        "strengths": [
            "산업 특화 멀티-Agent 오케스트레이션 플랫폼 (5종 전문Agent + 오케스트레이터)",
            "온톨로지/지식그래프 기반 제조 도메인 semantic grounding",
            "ReAct 기반 추론: 감지→추론→행동→관찰→판단 반복 루프",
            "HITL 승인 구조 기반 안전한 Agent 실행",
            "다국어(한국어/영어/일본어) 지원 AI Copilot",
        ],
    },
    "PdM": {
        "name": "예지보전",
        "desc": "다변량 센서 기반 상황 인지 및 PM 연계형 고장 예측 (정확도 90%+)",
        "keywords": ["예지보전", "고장예측", "RUL", "설비진단", "PHM", "보전"],
        "tech": "LSTM / Transformer / PHM / Edge Computing / PM 연동",
        "strengths": [
            "다변량 센서(진동/온도/전류) 기반 설비 상태 해석 및 Context 구성",
            "설비 잔여 수명(RUL) 예측 정확도 90%+",
            "PM 시스템 연동 자동 보전 일정 업데이트",
            "고장 원인·보전 조치·후속 영향 결과를 지식 베이스에 피드백하는 학습 구조",
        ],
    },
    "InfraDS": {
        "name": "데이터 인프라 / AAS 표준",
        "desc": "AAS(IEC 63278-1) 기반 제조 데이터 표준화 및 클라우드 인프라",
        "keywords": ["데이터", "인프라", "클라우드", "플랫폼", "MES", "ERP", "AAS"],
        "tech": "AAS / Catena-X / K8s / Kafka / MinIO / OPC-UA",
        "strengths": [
            "AAS(IEC 63278-1) 기반 제조 데이터 표준화 기술 보유",
            "Catena-X 기반 제조 데이터 스페이스 구축 경험",
            "OPC UA Tag ID ↔ AAS Asset ID 매핑 기반 표준 설비 모니터링",
            "제조 데이터 레이크 설계 및 ETL 파이프라인",
        ],
    },
}


# ═════════════════════════════════════════════════════════════════════════════
#  섹션 유형 판별 — 프롬프트 특화용
# ═════════════════════════════════════════════════════════════════════════════

def _classify_section(title: str) -> str:
    """섹션 제목을 유형으로 분류 → 특화 프롬프트 선택용."""
    t = title.lower()
    if any(kw in t for kw in ["개요", "목적", "배경", "정의", "범위"]):
        return "overview"
    if any(kw in t for kw in ["필요성", "현황", "동향", "시장", "기술수준"]):
        return "necessity"
    if any(kw in t for kw in ["목표", "내용", "개발내용", "세부목표", "구축내용"]):
        return "goals"
    if any(kw in t for kw in ["추진", "전략", "방법", "체계", "일정", "인력"]):
        return "strategy"
    if any(kw in t for kw in ["사업화", "시장진출", "매출", "수출"]):
        return "commercialization"
    if any(kw in t for kw in ["기대효과", "성과", "KPI", "효과", "활용"]):
        return "effects"
    if any(kw in t for kw in ["기업", "역량", "조직", "실적", "연구개발기관"]):
        return "company"
    if any(kw in t for kw in ["구축", "시스템", "설계", "아키텍처"]):
        return "goals"
    return "general"


# ═════════════════════════════════════════════════════════════════════════════
#  Step 1: 공고 본문에서 요구 섹션 구조 추출 (AI 기반)
# ═════════════════════════════════════════════════════════════════════════════

def extract_sections_from_notice(notice: Notice) -> List[Dict]:
    """공고 본문을 분석하여 사업계획서 요구 섹션 구조를 추출."""
    text = _prepare_notice_text(notice)
    if not text or len(text) < 50:
        return _default_sections(notice)

    try:
        from interx_engine.application.ports.gemini_port import generate, is_available
        if not is_available():
            return _default_sections(notice)
    except ImportError:
        return _default_sections(notice)

    system = (
        "당신은 한국 정부지원사업 전문가입니다. "
        "아래 공고 본문을 분석하여, 이 사업의 사업계획서에 들어가야 할 "
        "섹션 구조(목차)를 추출하세요.\n\n"
        "반드시 JSON 배열로만 응답하세요. 다른 텍스트 없이.\n"
        "형식:\n"
        '[{"title":"1. 섹션제목","subsections":["1.1 소제목","1.2 소제목"],"guidance":"이 섹션에서 요구하는 내용 설명"}]\n\n'
        "규칙:\n"
        "- 공고에서 명시적으로 요구하는 항목을 최우선으로 추출\n"
        "- 공고에 양식/목차가 없으면, 사업 성격에 맞는 표준 구조를 생성\n"
        "- 비용/예산/사업비 관련 섹션은 제외\n"
        "- 6~10개 대섹션이 적절 (R&D는 더 많을 수 있음)\n"
    )

    prompt = f"[공고 정보]\n제목: {notice.title}\n기관: {notice.agency or '-'}\n\n[본문]\n{text[:5000]}"

    try:
        raw = generate(prompt=prompt, system_instruction=system,
                       temperature=0.3, max_tokens=2048, timeout=40)
        sections = _parse_sections_json(raw)
        if sections:
            log.info("[BusinessPlan] 공고 분석 → %d개 섹션 추출", len(sections))
            return sections
    except Exception as e:
        log.warning("[BusinessPlan] 섹션 추출 실패: %s", e)

    return _default_sections(notice)


def extract_sections_from_template(template_text: str, notice_title: str = "") -> List[Dict]:
    """업로드된 양식 텍스트에서 섹션 구조를 추출."""
    if not template_text or len(template_text) < 30:
        return []

    try:
        from interx_engine.application.ports.gemini_port import generate, is_available
        if not is_available():
            return _parse_sections_regex(template_text)
    except ImportError:
        return _parse_sections_regex(template_text)

    system = (
        "당신은 한국 정부지원사업 사업계획서 양식 분석 전문가입니다. "
        "아래 텍스트는 사업계획서 양식입니다. "
        "이 양식의 섹션 구조(목차)를 추출하세요.\n\n"
        "반드시 JSON 배열로만 응답하세요.\n"
        '[{"title":"섹션제목","subsections":["소제목1","소제목2"],"guidance":"이 섹션에서 작성해야 할 내용"}]\n\n'
        "규칙:\n"
        "- 양식의 빈칸/작성란을 기반으로 구조 파악\n"
        "- 비용/예산/사업비 관련 섹션은 제외\n"
        "- 표지, 서약서, 인감 등 행정 양식은 제외\n"
    )

    prompt = f"[양식 제목] {notice_title}\n\n[양식 텍스트]\n{template_text[:5000]}"

    try:
        raw = generate(prompt=prompt, system_instruction=system,
                       temperature=0.2, max_tokens=2048, timeout=40)
        sections = _parse_sections_json(raw)
        if sections:
            log.info("[BusinessPlan] 양식 분석 → %d개 섹션 추출", len(sections))
            return sections
    except Exception as e:
        log.warning("[BusinessPlan] 양식 섹션 추출 실패: %s", e)

    return _parse_sections_regex(template_text)


# ═════════════════════════════════════════════════════════════════════════════
#  Step 2: 각 섹션 AI 내용 생성 (v3 — 지식베이스 기반)
# ═════════════════════════════════════════════════════════════════════════════

def generate_section_content(
    section: Dict,
    notice: Notice,
    solutions: List[str],
    company_name: str = "(주)인터엑스",
) -> str:
    """한 섹션의 내용을 Gemini로 생성 — 지식베이스 기반 상세 콘텐츠."""
    try:
        from interx_engine.application.ports.gemini_port import generate, is_available
        if not is_available():
            return _fallback_content(section, notice, solutions, company_name)
    except ImportError:
        return _fallback_content(section, notice, solutions, company_name)

    sec_type = _classify_section(section["title"])
    sol_info = _get_solutions_context(solutions)
    company_info = _get_company_context()
    guidance = section.get("guidance", "")

    # 섹션 유형별 추가 컨텍스트
    extra_context = ""
    if sec_type == "necessity":
        extra_context = f"\n{_get_market_context()}\n{_get_competitor_context()}\n{_get_standards_context()}"
    elif sec_type == "overview":
        extra_context = f"\n{_get_standards_context()}"
    elif sec_type == "commercialization":
        extra_context = f"\n{_get_market_context()}"
    elif sec_type == "company":
        extra_context = f"\n{_get_track_record_context()}"
    elif sec_type == "effects":
        extra_context = f"\n{_get_market_context()}"

    # 섹션 유형별 특화 지시사항
    type_instructions = _get_type_instructions(sec_type)

    system = (
        f"당신은 '{notice.title}' 사업의 사업계획서 작성 전문가입니다.\n"
        f"한국 정부지원사업(중기부/산업부/과기부 등) 사업계획서를 10년 이상 작성해 온 "
        f"최고 수준의 전문가로서, 실제 선정되는 수준의 상세하고 구체적인 내용을 작성합니다.\n\n"
        f"[기업 정보]\n{company_info}\n\n"
        f"[작성 규칙]\n"
        "1. 한국어, 공식 문서 어투 (경어체 '-임', '-함' 종결)\n"
        "2. 비용/예산/금액 정보는 절대 포함하지 마세요\n"
        "3. [작성 필요] 같은 플레이스홀더 절대 사용 금지\n"
        "4. 뻔한 일반론('생산성 향상', '품질 개선') 대신 구체적 기술명, 방법론, 수치를 기술\n"
        "5. 번호 매기기(○, -, □), 불릿 활용하여 가독성 확보\n"
        "6. 해당 섹션에 적합한 내용만 작성 (섹션 범위를 벗어나지 말 것)\n"
        "7. 구체적인 기술명(AAS, FMEA, OntoRef, SPC, RMS 등)과 표준명(IEC 63278 등)을 적극 활용\n"
        "8. 1500자 이상 상세하게 작성\n\n"
        f"{type_instructions}\n"
    )

    subsec_str = ""
    if section.get("subsections"):
        subsec_str = "\n소제목:\n" + "\n".join(f"  - {s}" for s in section["subsections"])

    prompt = (
        f"[공고 정보]\n"
        f"공고명: {notice.title}\n"
        f"주관기관: {notice.agency or '-'}\n"
        f"사업 요약: {(notice.summary or notice.body_text[:500])[:500]}\n\n"
        f"[기업 솔루션 상세]\n{sol_info}\n\n"
        f"{extra_context}\n\n"
        f"[작성 요청]\n"
        f"섹션: {section['title']}\n"
        f"{subsec_str}\n"
        f"{'작성 지침: ' + guidance if guidance else ''}\n\n"
        f"위 정보를 모두 활용하여 이 섹션의 내용을 1500자 이상 상세하게 작성하세요. "
        f"구체적인 기술명, 시스템명, 방법론을 반드시 포함하세요."
    )

    try:
        result = generate(prompt=prompt, system_instruction=system,
                          temperature=0.5, max_tokens=4096, timeout=60)
        if result and len(result) > 100:
            return result
    except Exception as e:
        log.warning("[BusinessPlan] 섹션 생성 실패 (%s): %s", section["title"], e)

    return _fallback_content(section, notice, solutions, company_name)


def _get_type_instructions(sec_type: str) -> str:
    """섹션 유형별 특화 작성 지시사항."""
    instructions = {
        "overview": (
            "[개요 섹션 작성 패턴]\n"
            "1. (기존 공정) 현재 제조 현장의 한계점을 구체적으로 기술\n"
            "   예: '설비, 공정, 검사, 품질, 생산계획 영역이 PM·MES·SPC·RMS 등 개별 Legacy 시스템과 "
            "작업자 경험 중심으로 분절 운영'\n"
            "2. (개발 기술) 해결하고자 하는 기술의 구체적 구성 요소 나열\n"
            "   예: '산업 특화 멀티-Agent를 개발. 전문 Agent(생산관리, 예지보전, 품질 분석, 공정자동화, "
            "생산계획)와 오케스트레이션 Agent, 지식 그래프 기반 체계로 구성'\n"
            "3. 연구개발 기술의 범위 (세부 기술 목록 나열)\n"
        ),
        "necessity": (
            "[필요성/현황 섹션 작성 패턴]\n"
            "1. 국내 기술 동향 — 현재 수준과 한계(Level 2~3)를 구체적으로 기술\n"
            "2. 국내 시장 규모 — 제공된 시장 데이터를 반드시 인용 (연도, 금액, CAGR)\n"
            "3. 국내 경쟁기관 분석 — 각 경쟁사의 제품명과 특징, 인터엑스와의 차이점\n"
            "4. 국외 기술 동향 — 글로벌 선도 기업(Siemens, Microsoft 등)의 최신 기술\n"
            "5. 국외 시장 규모 — 글로벌 시장 데이터 인용\n"
            "6. 인터엑스 기술 수준 — 보유 기술과 Level 4 달성 가능성\n"
            "주의: 반드시 구체적 기업명, 제품명, 수치를 포함할 것\n"
        ),
        "goals": (
            "[목표/내용 섹션 작성 패턴]\n"
            "1. 최종 목표를 2~3문장으로 명확히 기술\n"
            "2. 세부 목표를 [세부 목표 1], [세부 목표 2]... 형식으로 나열\n"
            "3. 각 세부 목표에 3개 이상의 구체적 기술 개발 항목 포함\n"
            "   예: '(지식 구조화 기술) 설비, 공정, 검사, 품질, 생산계획, 보전 이력 등 이종 제조 데이터를 "
            "통합 관리할 수 있도록 산업현장 데이터 항목과 관계 구조를 정의'\n"
            "4. 각 기술 항목에 구체적 방법론 명시 (4M1E, FMEA, RPN, AAS 등)\n"
        ),
        "strategy": (
            "[추진전략/체계 섹션 작성 패턴]\n"
            "1. 추진 전략 — 단계별 접근법 (1단계: 기반구축 → 2단계: 개발 → 3단계: 실증)\n"
            "2. 추진 체계 — 총괄PM, 기술개발팀, 수요기업 역할 분담\n"
            "3. 기술개발팀 편성 — 팀별 역할 (AI팀, DT팀, 시스템연계팀 등)\n"
            "4. 품질 관리 — HITL 승인 구조, 코드 리뷰, 테스트 자동화\n"
        ),
        "commercialization": (
            "[사업화 전략 섹션 작성 패턴]\n"
            "1. 목표 시장 — 제공된 시장 규모 데이터 인용\n"
            "2. 사업화 단계 — 실증 → 상용화 → 확산\n"
            "3. 매출 전망 — 연차별 목표 (구체적 수치)\n"
            "4. 시장 진출 전략 — B2B SaaS, 컨소시엄, 해외 진출\n"
            "5. 고용 창출 효과\n"
        ),
        "effects": (
            "[기대효과 섹션 작성 패턴]\n"
            "1. 정량적 기대효과 — KPI별 목표 수치 (불량률 30%↓, 설비가동률 15%↑ 등)\n"
            "2. 정성적 기대효과 — 기술 자립, 표준화 기여 등\n"
            "3. 성과 활용방안 — 특허, 논문, 기술이전, 표준화 기여\n"
            "4. 산업 파급효과 — 관련 산업 생태계 활성화\n"
        ),
        "company": (
            "[기업 역량 섹션 작성 패턴]\n"
            "1. 기업 개요 — 설립, 대표, 인력, 소재지\n"
            "2. 기술 역량 — 보유 기술(AAS, LLM, Agent 등) 상세 설명\n"
            "3. 수행 실적 — 유사 과제 수행 경험\n"
            "4. 지식재산권 — 관련 특허 보유 현황\n"
            "5. 인증 현황 — 벤처기업, 이노비즈, 기업부설연구소\n"
        ),
    }
    return instructions.get(sec_type, (
        "[일반 섹션 작성 규칙]\n"
        "- 구체적 기술명과 방법론을 포함\n"
        "- 인터엑스의 관련 역량을 자연스럽게 연결\n"
        "- 뻔한 일반론 대신 실무적 내용 기술\n"
    ))


# ═════════════════════════════════════════════════════════════════════════════
#  Step 3: DOCX 문서 조립
# ═════════════════════════════════════════════════════════════════════════════

def build_docx(
    notice: Notice,
    sections: List[Dict],
    section_contents: Dict[str, str],
    score_card: Optional[ScoreCard] = None,
    solutions: List[str] = None,
    company_name: str = "(주)인터엑스",
) -> Optional[str]:
    """섹션 구조 + 생성된 내용 → DOCX 파일 조립."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        log.error("[BusinessPlan] python-docx 미설치")
        return None

    doc = Document()

    # ── 기본 스타일 ───────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    # ── 표지 ──────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(notice.title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0A, 0x16, 0x28)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("사 업 계 획 서")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x4E, 0x92)

    for _ in range(4):
        doc.add_paragraph()

    for line in [
        f"신 청 기 업 : {company_name}",
        f"작 성 일 자 : {date.today().isoformat()}",
        f"공 고 기 관 : {notice.agency or notice.ministry or '-'}",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(13)

    doc.add_page_break()

    # ── 목차 ──────────────────────────────────────────────────────────────
    doc.add_heading("목 차", level=1)
    for sec in sections:
        doc.add_paragraph(sec["title"], style="List Number")
        for sub in sec.get("subsections", []):
            p = doc.add_paragraph(f"    {sub}")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ── 공고 요약 정보 테이블 ─────────────────────────────────────────────
    doc.add_heading("공고 요약 정보", level=1)
    info_rows = [
        ("공고명", notice.title),
        ("주관기관", notice.agency or notice.ministry or "-"),
        ("마감일", notice.deadline_date or "-"),
        ("사업유형", notice.category or "-"),
        ("출처", notice.site),
    ]
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.style = "Light Shading Accent 1"
    for i, (k, v) in enumerate(info_rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # 적합도 정보
    if score_card:
        doc.add_paragraph()
        grade = score_card.priority_grade or "-"
        p = doc.add_paragraph()
        run = p.add_run(f"InterX 적합도: {grade} 등급 | "
                        f"적합도 {score_card.fitness_score:.0f}점 | "
                        f"우선순위 {score_card.priority_score:.0f}점")
        run.bold = True
        run.font.size = Pt(11)

    doc.add_page_break()

    # ── 본문 섹션 ─────────────────────────────────────────────────────────
    for sec in sections:
        doc.add_heading(sec["title"], level=1)

        content = section_contents.get(sec["title"], "")

        if not content:
            doc.add_paragraph(f"[작성 필요: {sec['title']}]")
            continue

        # 소제목이 있으면 heading level 2로
        for sub in sec.get("subsections", []):
            pass

        # 내용 파싱 및 추가
        _add_content_to_doc(doc, content)

        doc.add_paragraph()  # 섹션 간 여백

    # ── 부록: 공고 원문 ───────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("부록: 공고 원문 참조", level=1)
    url = notice.detail_url or notice.notice_link or ""
    if url:
        doc.add_paragraph(f"공고 링크: {url}")
    body = notice.body_text or notice.summary or ""
    if body:
        excerpt = body[:2500]
        for line in excerpt.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    # ── 면책 ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "본 문서는 AI가 자동 생성한 사업계획서 초안입니다. "
        "실제 제출 전 반드시 검토/수정하세요. "
        "비용/예산 정보는 의도적으로 제외되었습니다."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # ── 저장 ──────────────────────────────────────────────────────────────
    import tempfile
    out_dir = Path(tempfile.gettempdir()) / "interx_business_plans"
    out_dir.mkdir(parents=True, exist_ok=True)

    safe = "".join(c for c in notice.title[:35] if c.isalnum() or c in " _-").strip()
    safe = safe or notice.notice_id
    fname = out_dir / f"사업계획서_{safe}.docx"
    doc.save(str(fname))
    log.info("[BusinessPlan] 저장: %s", fname.name)
    return str(fname)


# ═════════════════════════════════════════════════════════════════════════════
#  통합 실행 함수
# ═════════════════════════════════════════════════════════════════════════════

def generate_business_plan(
    notice: Notice,
    score_card: Optional[ScoreCard] = None,
    template_text: str = "",
    company_name: str = "(주)인터엑스",
    progress_callback=None,
) -> Optional[str]:
    """공고 맞춤형 사업계획서 생성 (통합 진입점)."""
    def _progress(pct, msg):
        if progress_callback:
            progress_callback(pct, msg)
        log.info("[BusinessPlan] [%d%%] %s", pct, msg)

    # 지식베이스 사전 로드
    _progress(3, "지식베이스 로딩...")
    _load_knowledge_base()

    _progress(5, "솔루션 매칭 중...")
    solutions = _detect_relevant_solutions(notice, score_card)

    # Step 1: 섹션 구조 추출
    _progress(10, "섹션 구조 분석 중...")
    if template_text:
        sections = extract_sections_from_template(template_text, notice.title)
        _progress(25, f"양식 분석 완료 → {len(sections)}개 섹션")
    else:
        sections = extract_sections_from_notice(notice)
        _progress(25, f"공고 분석 완료 → {len(sections)}개 섹션")

    if not sections:
        _progress(25, "기본 섹션 구조 사용")
        sections = _default_sections(notice)

    # Step 2: 각 섹션 내용 생성
    section_contents = {}
    total = len(sections)
    for i, sec in enumerate(sections):
        pct = 30 + int(60 * i / max(total, 1))
        _progress(pct, f"섹션 생성 중: {sec['title']} ({i+1}/{total})")
        content = generate_section_content(sec, notice, solutions, company_name)
        section_contents[sec["title"]] = content

    # Step 3: DOCX 조립
    _progress(92, "DOCX 문서 조립 중...")
    path = build_docx(
        notice=notice,
        sections=sections,
        section_contents=section_contents,
        score_card=score_card,
        solutions=solutions,
        company_name=company_name,
    )

    _progress(100, "완료!")
    return path


# ═════════════════════════════════════════════════════════════════════════════
#  유틸리티 함수들
# ═════════════════════════════════════════════════════════════════════════════

def _prepare_notice_text(notice: Notice) -> str:
    """공고 분석용 텍스트 준비."""
    parts = []
    if notice.summary:
        parts.append(notice.summary)
    if notice.body_text:
        parts.append(notice.body_text[:5000])
    if notice.structured:
        for k, v in notice.structured.items():
            if v:
                parts.append(f"{k}: {v[:500]}")
    return "\n".join(parts)


def _detect_relevant_solutions(notice: Notice, sc: Optional[ScoreCard]) -> List[str]:
    """공고에 적합한 솔루션 Top 3."""
    text = f"{notice.title} {notice.summary} {notice.body_text[:2000]}".lower()

    # ScoreCard 솔루션 점수 우선
    if sc and sc.solution_scores:
        ranked = sorted(
            [(k, v) for k, v in sc.solution_scores.items() if v > 0],
            key=lambda x: -x[1],
        )
        if ranked:
            return [k for k, _ in ranked[:3]]

    # 키워드 매칭 fallback
    scores = {}
    for sol_key, sol_info in INTERX_CAPABILITIES.items():
        scores[sol_key] = sum(1 for kw in sol_info["keywords"] if kw.lower() in text)
    ranked = sorted(scores.items(), key=lambda x: -x[1])
    return [k for k, v in ranked[:3] if v > 0] or ["QualityAI", "PdM"]


def _solutions_text(solutions: List[str]) -> str:
    """솔루션 정보를 텍스트로 정리 (간략 버전)."""
    lines = []
    for s in solutions:
        cap = INTERX_CAPABILITIES.get(s, {})
        if cap:
            lines.append(f"- {cap['name']}: {cap['desc']}")
            lines.append(f"  기술스택: {cap['tech']}")
            for st_item in cap.get("strengths", [])[:2]:
                lines.append(f"  역량: {st_item}")
    return "\n".join(lines) if lines else "범용 제조 AI 솔루션"


def _parse_sections_json(raw: str) -> List[Dict]:
    """Gemini 응답에서 JSON 배열 추출."""
    import json

    # ```json ... ``` 블록 추출
    m = re.search(r"```(?:json)?\s*(\[.*?\])\s*```", raw, re.DOTALL)
    if m:
        raw = m.group(1)
    else:
        # 순수 JSON 배열 추출
        m = re.search(r"\[.*\]", raw, re.DOTALL)
        if m:
            raw = m.group(0)

    try:
        data = json.loads(raw)
        if isinstance(data, list):
            result = []
            for item in data:
                if isinstance(item, dict) and "title" in item:
                    result.append({
                        "title": item["title"],
                        "subsections": item.get("subsections", []),
                        "guidance": item.get("guidance", ""),
                    })
            return result
    except (json.JSONDecodeError, TypeError):
        pass
    return []


def _parse_sections_regex(text: str) -> List[Dict]:
    """정규식으로 양식 텍스트에서 번호 매겨진 섹션 제목 추출."""
    patterns = [
        r"^(\d+[\.\)]\s+.+)$",
        r"^([IⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+[\.\)]\s+.+)$",
        r"^([가나다라마바사아자차카타파하][\.\)]\s+.+)$",
    ]
    sections = []
    for line in text.split("\n"):
        line = line.strip()
        for pat in patterns:
            if re.match(pat, line):
                if any(kw in line for kw in ["사업비", "예산", "비용", "재원", "출연금"]):
                    continue
                sections.append({
                    "title": line,
                    "subsections": [],
                    "guidance": "",
                })
                break
    return sections


def _default_sections(notice: Notice) -> List[Dict]:
    """공고 분석 실패 시 기본 섹션 (사업 성격 추론)."""
    title_lower = f"{notice.title} {notice.summary}".lower()

    # 스마트공장 계열
    if any(kw in title_lower for kw in ["스마트공장", "스마트 공장", "제조ai", "스마트팩토리"]):
        return [
            {"title": "1. 스마트공장 구축 개요", "subsections": ["1.1 구축 목표", "1.2 현재 수준 진단", "1.3 목표 수준"], "guidance": ""},
            {"title": "2. 기업 현황", "subsections": ["2.1 기업 개요", "2.2 주요 생산품 및 공정", "2.3 기존 정보화 현황"], "guidance": ""},
            {"title": "3. 기술 현황 및 필요성", "subsections": ["3.1 국내외 기술 동향", "3.2 시장 현황", "3.3 경쟁사 분석"], "guidance": ""},
            {"title": "4. 구축 내용", "subsections": ["4.1 핵심 기술 적용 방안", "4.2 시스템 구성", "4.3 데이터 수집/활용 계획"], "guidance": ""},
            {"title": "5. 추진 체계 및 일정", "subsections": ["5.1 추진 체계", "5.2 단계별 일정", "5.3 투입 인력"], "guidance": ""},
            {"title": "6. 기대효과 및 성과지표", "subsections": ["6.1 정량적 기대효과", "6.2 정성적 기대효과", "6.3 KPI"], "guidance": ""},
        ]

    # R&D 계열
    if any(kw in title_lower for kw in ["연구개발", "r&d", "기술개발", "연구과제", "에이전트"]):
        return [
            {"title": "1. 필요성 및 현황", "subsections": ["1.1 개발 대상 기술/제품 개요", "1.2 국내외 기술과 시장 현황", "1.3 관련 지식재산권/표준화/인증 현황"], "guidance": ""},
            {"title": "2. 과제 목표 및 내용", "subsections": ["2.1 최종목표", "2.2 세부 목표별 개발내용", "2.3 수행일정"], "guidance": ""},
            {"title": "3. 추진전략 및 추진체계", "subsections": ["3.1 기술개발 추진전략/방법", "3.2 추진체계", "3.3 기술개발팀 편성"], "guidance": ""},
            {"title": "4. 성과 활용방안 및 기대효과", "subsections": ["4.1 성과 활용방안", "4.2 기대효과", "4.3 성과 기술기여도"], "guidance": ""},
            {"title": "5. 사업화 전략 및 계획", "subsections": ["5.1 사업화 전략", "5.2 시장 진출 계획", "5.3 기대효과"], "guidance": ""},
            {"title": "6. 연구개발기관 현황", "subsections": ["6.1 연구책임자 현황", "6.2 연구개발 실적", "6.3 연구시설/장비 보유현황"], "guidance": ""},
        ]

    # 바우처 계열
    if any(kw in title_lower for kw in ["바우처", "voucher", "ai솔루션"]):
        return [
            {"title": "1. 사업 개요", "subsections": ["1.1 사업 목적", "1.2 AI 도입 필요성", "1.3 현재 업무 프로세스"], "guidance": ""},
            {"title": "2. AI 솔루션 적용 방안", "subsections": ["2.1 적용 기술", "2.2 시스템 구성", "2.3 데이터 활용 계획"], "guidance": ""},
            {"title": "3. 기업 현황 및 역량", "subsections": ["3.1 수요기업 현황", "3.2 공급기업 역량", "3.3 수행실적"], "guidance": ""},
            {"title": "4. 추진 일정 및 체계", "subsections": ["4.1 추진 일정", "4.2 역할 분담"], "guidance": ""},
            {"title": "5. 기대효과", "subsections": ["5.1 정량적 효과", "5.2 정성적 효과"], "guidance": ""},
        ]

    # 일반
    return [
        {"title": "1. 사업 개요", "subsections": ["1.1 사업 목적", "1.2 사업 범위", "1.3 추진 필요성"], "guidance": ""},
        {"title": "2. 기술 현황 및 필요성", "subsections": ["2.1 국내외 기술 동향", "2.2 시장 현황", "2.3 경쟁사 분석"], "guidance": ""},
        {"title": "3. 기업 현황 및 역량", "subsections": ["3.1 기업 개요", "3.2 기술 역량", "3.3 수행 실적"], "guidance": ""},
        {"title": "4. 기술 개발 내용", "subsections": ["4.1 핵심 기술", "4.2 적용 방안", "4.3 시스템 구성"], "guidance": ""},
        {"title": "5. 추진 체계 및 일정", "subsections": ["5.1 추진 체계", "5.2 추진 일정"], "guidance": ""},
        {"title": "6. 기대효과", "subsections": ["6.1 정량적 효과", "6.2 정성적 효과", "6.3 성과지표"], "guidance": ""},
    ]


def _fallback_content(
    section: Dict,
    notice: Notice,
    solutions: List[str],
    company_name: str,
) -> str:
    """Gemini 실패 시 지식베이스 기반 풍부한 fallback 콘텐츠."""
    kb = _load_knowledge_base()
    company = kb.get("company", {})
    sol_data = kb.get("solutions", {})
    market = kb.get("market_data", {})
    comp = kb.get("competitors", {})

    sol_names = [INTERX_CAPABILITIES.get(s, {}).get("name", s) for s in solutions]
    title = section["title"].lower()

    # ── 개요 / 목적 / 배경 / 필요성 ──────────────────────────────────────
    if any(kw in title for kw in ["개요", "목적", "배경", "필요성", "현황"]):
        lines = [
            f"본 사업은 '{notice.title}'에 따라 {company_name}가 보유한 핵심 기술을 활용하여 추진하는 과제임.\n",
            "■ 기존 공정의 한계",
            "현재 제조 현장에서는 설비, 공정, 검사, 품질, 생산계획 영역이 PM(Preventive Maintenance), "
            "MES(Manufacturing Execution System), SPC(Statistical Process Control), "
            "RMS(Recipe Management System) 등 개별 Legacy 시스템과 작업자 경험 중심으로 분절 운영되어, "
            "고장 예측 결과가 생산계획에 즉시 반영되지 못하고, 검사 결과와 원인 공정-Recipe 조정-생산량 조정 간 "
            "연결이 미흡한 상태임.\n",
            "■ 개발 기술",
            f"{company_name}는 기구축된 데이터/모델 자산을 활용하여, 설비-공정-검사-품질-생산계획 데이터를 "
            "통합 해석하고 현장 시스템과 연계하는 산업 특화 멀티-Agent를 개발하고자 함.\n",
            "전문 Agent 구성:",
        ]
        # 에이전트 정보 추가
        ma = sol_data.get("multi_agent_platform", {})
        for agent in ma.get("agents", [])[:5]:
            lines.append(f"  - {agent['name']}: {agent['role']}")

        lines.append("\n■ 인터엑스 핵심 역량")
        for cap in company.get("core_competency", [])[:5]:
            lines.append(f"  - {cap}")

        # 시장 데이터
        dt_kr = market.get("digital_twin", {}).get("korea", {})
        if dt_kr:
            lines.append(f"\n■ 시장 규모")
            lines.append(f"  - 한국 디지털트윈 시장: {dt_kr.get('size_2024','')} (2024) → {dt_kr.get('size_2033','')} (2033), CAGR {dt_kr.get('cagr','')}")

        ai_g = market.get("ai_agent", {}).get("global", {})
        if ai_g:
            lines.append(f"  - 글로벌 AI Agent 시장: {ai_g.get('size_2024','')} (2024) → {ai_g.get('size_2030','')} (2030), CAGR {ai_g.get('cagr','')}")

        # 경쟁사
        if comp.get("domestic"):
            lines.append("\n■ 국내 주요 경쟁기관")
            for c in comp.get("domestic", [])[:4]:
                lines.append(f"  - {c['name']}: {c['product']} — {c['focus']}")
            diff = comp.get("interx_differentiation", "")
            if diff:
                lines.append(f"\n■ 인터엑스 차별점\n  {diff.strip()[:400]}")

        return "\n".join(lines)

    # ── 기업 역량 ─────────────────────────────────────────────────────────
    if any(kw in title for kw in ["기업", "역량", "조직", "실적", "연구개발기관"]):
        lines = [
            f"■ 기업 개요",
            f"{company_name}은 {company.get('established',2019)}년 설립된 {company.get('type','제조 AI 전문기업')}으로, "
            f"대표이사 {company.get('ceo','박정윤')}, 소재지 {company.get('location','대전광역시')}, "
            f"인력 {company.get('employees',25)}명 규모의 기업임.\n",
            f"■ 핵심 기술 역량",
            f"{company.get('description','').strip()}\n",
        ]
        for cap in company.get("core_competency", []):
            lines.append(f"  - {cap}")

        lines.append("\n■ 주요 수행실적")
        for p in kb.get("track_record", {}).get("projects", []):
            lines.append(f"  - [{p['name']}] {p.get('description','')}")

        lines.append("\n■ 지식재산권")
        for pat in company.get("patents", []):
            lines.append(f"  - {pat}")

        lines.append("\n■ 인증 현황")
        for cert in company.get("certifications", []):
            lines.append(f"  - {cert}")

        return "\n".join(lines)

    # ── 기술 / 구축 / 개발 ────────────────────────────────────────────────
    if any(kw in title for kw in ["기술", "구축", "개발", "내용", "목표"]):
        lines = ["■ 핵심 기술 적용 방안\n"]
        for s, n in zip(solutions, sol_names):
            cap = INTERX_CAPABILITIES.get(s, {})
            lines.append(f"□ {n}")
            lines.append(f"  {cap.get('desc', '')}")
            lines.append(f"  기술스택: {cap.get('tech', '')}")
            for st in cap.get("strengths", []):
                lines.append(f"  - {st}")
            lines.append("")

        # 상세 솔루션 정보
        ont = sol_data.get("ontology_knowledge", {})
        if ont:
            lines.append(f"□ {ont.get('name','온톨로지 기반 지식체계')}")
            for feat in ont.get("key_features", []):
                lines.append(f"  - {feat}")

        lines.append("\n■ 시스템 아키텍처")
        arch = sol_data.get("multi_agent_platform", {}).get("architecture", "")
        if arch:
            lines.append(f"  {arch.strip()[:500]}")

        safety = sol_data.get("multi_agent_platform", {}).get("safety", "")
        if safety:
            lines.append(f"\n■ 안전성 확보 방안")
            lines.append(f"  {safety.strip()[:300]}")

        return "\n".join(lines)

    # ── 추진 체계 ─────────────────────────────────────────────────────────
    if "추진" in title:
        return (
            f"■ 추진 체계\n"
            f"  - 총괄 PM: {company_name} 대표이사 {company.get('ceo','박정윤')}\n"
            f"  - 기술개발: AI/데이터 팀 (산업 AI Agent, 온톨로지, DT 담당)\n"
            f"  - 시스템연계: 인프라팀 (PM/MES/SPC/RMS/APS 연동)\n"
            f"  - 현장 실증: 수요기업 협력 (데이터 제공, 현장 검증)\n\n"
            f"■ 추진 전략\n"
            f"  1단계 (기반 구축): 온톨로지/지식그래프 구축, 데이터 파이프라인 설계\n"
            f"  2단계 (핵심 개발): 멀티-Agent 플랫폼 개발, 개별 Agent 기능 구현\n"
            f"  3단계 (통합/실증): 시스템 통합, 현장 실증, 성능 검증\n\n"
            f"■ 품질 관리\n"
            f"  - HITL(Human-in-the-loop) 승인 구조 적용\n"
            f"  - Agent 실행 이력, 승인/반려 이력, 재시도 이력 누적 관리\n"
            f"  - 코드 리뷰 및 테스트 자동화\n"
        )

    # ── 기대효과 / 성과 ───────────────────────────────────────────────────
    if any(kw in title for kw in ["기대", "성과", "KPI", "효과", "활용"]):
        lines = [
            "■ 정량적 기대효과",
            "  - 설비 가동률: 15% 이상 향상",
            "  - 제품 불량률: 30% 이상 감소",
            "  - 설비 고장 예측 정확도: 90% 이상",
            "  - 생산계획 수립 시간: 50% 단축 (12h → 4h 이내)",
            "  - 정비 비용: 30% 이상 절감",
            "  - 생산성: 20% 이상 향상\n",
            "■ 정성적 기대효과",
            "  - 데이터 기반 의사결정 체계 구축으로 작업자 경험 의존도 탈피",
            "  - AAS(IEC 63278-1) 표준 기반 제조 데이터 상호운용성 확보",
            "  - 멀티-Agent 기반 자율 제어(Level 4) 구현으로 제조 AI 기술 선도",
            "  - 온톨로지/FMEA 기반 지식 체계 구축으로 축적형 제조 지능화 실현\n",
            "■ 성과 활용방안",
            "  - 과제 성과를 기반으로 유사 제조 현장(반도체/자동차/화학) 확산",
            "  - 솔루션 패키지화를 통한 B2B SaaS 모델 구축",
            "  - 관련 특허 출원 및 기술 표준화 기여",
        ]

        # 시장 데이터 추가
        dt_kr = market.get("digital_twin", {}).get("korea", {})
        if dt_kr:
            lines.append(f"\n■ 목표 시장 규모")
            lines.append(f"  한국 디지털트윈 시장: {dt_kr.get('size_2024','')} → {dt_kr.get('size_2033','')} (CAGR {dt_kr.get('cagr','')})")

        return "\n".join(lines)

    # ── 사업화 ────────────────────────────────────────────────────────────
    if any(kw in title for kw in ["사업화", "시장", "매출"]):
        lines = [
            "■ 사업화 전략",
            f"  {company_name}는 본 과제의 성과를 기반으로 다음과 같은 사업화 전략을 추진함.\n",
            "  1) 실증 → 상용화 → 확산 단계적 접근",
            "    - 1단계: 수요기업 현장 실증 및 성과 검증",
            "    - 2단계: 솔루션 패키지화 및 B2B SaaS 모델 구축",
            "    - 3단계: 유사 산업(자동차, 화학, 소재) 확산 및 해외 진출\n",
        ]

        dt_kr = market.get("digital_twin", {}).get("korea", {})
        ai_g = market.get("ai_agent", {}).get("global", {})
        if dt_kr or ai_g:
            lines.append("■ 목표 시장")
            if dt_kr:
                lines.append(f"  - 한국 DT 시장: {dt_kr.get('size_2024','')} (2024) → {dt_kr.get('size_2033','')} (2033)")
            if ai_g:
                lines.append(f"  - 글로벌 AI Agent 시장: {ai_g.get('size_2024','')} (2024) → {ai_g.get('size_2030','')} (2030)")

        kai = market.get("korea_ai_industry", {})
        if kai:
            lines.append(f"\n■ 한국 AI 산업 현황")
            lines.append(f"  총매출 {kai.get('total_revenue_2024','')}, B2B 비중 {kai.get('b2b_share','')}")

        diff = comp.get("interx_differentiation", "")
        if diff:
            lines.append(f"\n■ 시장 차별화 전략\n  {diff.strip()[:400]}")

        lines.append("\n■ 매출 전망")
        lines.append("  - 1년차: 기반 구축 및 실증 (파일럿 매출)")
        lines.append("  - 2년차: B2B 솔루션 론칭 및 초기 고객 확보")
        lines.append("  - 3년차: 확산 및 해외 진출")

        return "\n".join(lines)

    # ── 그 외 ─────────────────────────────────────────────────────────────
    lines = [f"■ {section['title']}\n"]
    lines.append(f"{company_name}는 {', '.join(sol_names)} 기술을 보유한 제조 AI 전문기업으로, "
                 f"'{notice.title}' 사업에 다음과 같이 참여하고자 함.\n")
    for s, n in zip(solutions, sol_names):
        cap = INTERX_CAPABILITIES.get(s, {})
        lines.append(f"  - {n}: {cap.get('desc', '')}")
        for st in cap.get("strengths", [])[:2]:
            lines.append(f"    {st}")
    return "\n".join(lines)


def _add_content_to_doc(doc, content: str):
    """텍스트 내용을 DOCX 문단으로 변환."""
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        # 마크다운 볼드 → 일반 텍스트
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        # 마크다운 헤더 → 굵은 텍스트
        header_match = re.match(r"^#{1,3}\s+(.+)", line)
        if header_match:
            from docx.shared import Pt
            p = doc.add_paragraph()
            run = p.add_run(header_match.group(1))
            run.bold = True
            run.font.size = Pt(11)
            continue
        # ■ □ 헤더 → 굵은 텍스트
        if re.match(r"^[■□]\s", line):
            from docx.shared import Pt
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(10.5)
            continue
        # 불릿 아이템
        if re.match(r"^[-*]\s", line):
            doc.add_paragraph(line.lstrip("-* "), style="List Bullet")
        elif re.match(r"^\d+[\.\)]\s", line):
            doc.add_paragraph(line, style="List Number")
        elif re.match(r"^[○●]\s", line):
            doc.add_paragraph(line.lstrip("○● "), style="List Bullet")
        else:
            doc.add_paragraph(line)


def parse_uploaded_file(file_bytes: bytes, filename: str) -> str:
    """업로드된 파일(PDF/HWP/HWPX/TXT)에서 텍스트 추출."""
    ext = Path(filename).suffix.lower()

    if ext == ".txt":
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("cp949", errors="replace")

    if ext == ".pdf":
        return _extract_pdf_text(file_bytes)

    if ext == ".hwpx":
        return _extract_hwpx_text(file_bytes)

    if ext == ".hwp":
        return _extract_hwp_text(file_bytes)

    log.warning("[BusinessPlan] 미지원 파일 형식: %s", ext)
    return ""


def _extract_pdf_text(data: bytes) -> str:
    """PDF 텍스트 추출."""
    try:
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(data))
        texts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
        return "\n".join(texts)
    except ImportError:
        log.warning("[BusinessPlan] pypdf 미설치")
    except Exception as e:
        log.warning("[BusinessPlan] PDF 파싱 실패: %s", e)
    return ""


def _extract_hwpx_text(data: bytes) -> str:
    """HWPX (ZIP+XML) 텍스트 추출."""
    import zipfile
    import io
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            texts = []
            for name in sorted(zf.namelist()):
                if name.startswith("Contents/section") and name.endswith(".xml"):
                    try:
                        import xml.etree.ElementTree as ET
                        tree = ET.parse(zf.open(name))
                        root = tree.getroot()
                        for elem in root.iter():
                            if elem.text and elem.text.strip():
                                texts.append(elem.text.strip())
                    except Exception:
                        pass
            return "\n".join(texts)
    except Exception as e:
        log.warning("[BusinessPlan] HWPX 파싱 실패: %s", e)
    return ""


def _extract_hwp_text(data: bytes) -> str:
    """HWP (바이너리 OLE) 텍스트 추출 — 노이즈 필터링 포함."""
    try:
        import olefile
        import zlib
        import io

        ole = olefile.OleFileIO(io.BytesIO(data))
        texts = []
        for entry in ole.listdir():
            name = "/".join(entry)
            if name.startswith("BodyText/Section"):
                raw = ole.openstream(name).read()
                try:
                    decompressed = zlib.decompress(raw, -15)
                except zlib.error:
                    decompressed = raw
                try:
                    t = decompressed.decode("utf-16-le", errors="replace")
                except Exception:
                    t = decompressed.decode("cp949", errors="replace")
                clean = re.findall(
                    r'[가-힣ㄱ-ㆎa-zA-Z0-9 \-~\.\,\(\)\[\]\/\:\;\!\?\%\+\=\@\#\&\*\"\']+',
                    t,
                )
                clean_text = " ".join(c.strip() for c in clean if len(c.strip()) > 1)
                if clean_text:
                    texts.append(clean_text)
        ole.close()
        return "\n".join(texts)
    except ImportError:
        log.warning("[BusinessPlan] olefile 미설치")
    except Exception as e:
        log.warning("[BusinessPlan] HWP 파싱 실패: %s", e)
    return ""
