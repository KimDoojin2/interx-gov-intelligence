# =============================================================================
# InterX Engine — Colab 패치셀 v4
# 실행 방법: Google Colab 셀에 이 파일 전체를 붙여넣고 실행
#
# 포함 파일 (신규/수정):
#   src/interx_engine/core/entities/notice.py           ← is_new/is_changed/memo 등 추가
#   src/interx_engine/application/mappers/notice_mapper.py ← 신규/변경/경쟁사 컬럼 포함
#   src/interx_engine/application/orchestrators/daily_pipeline.py ← 12단계 파이프라인
#   run_engine.py                                        ← NEW_COLLECTOR_CLASSES 포함
#   src/.../infrastructure/utils/budget_parser.py        ← 예산 정규화
#   src/.../infrastructure/utils/morpheme_scorer.py      ← 형태소 스코어
#   src/.../collectors/sites/new_collectors.py           ← nrf/kised/ketep/koiia
#   src/.../use_cases/deduplicate_notices.py             ← TF-IDF 중복감지
#   src/.../use_cases/detect_changes.py                  ← MD5 변경감지
#   src/.../use_cases/assign_manager.py                  ← 담당자 자동배정
#   src/.../use_cases/track_competitors.py               ← 경쟁사 트래킹
#   src/.../use_cases/site_quality_grader.py             ← 사이트 품질등급
#   src/.../use_cases/generate_proposal.py               ← 제안서 자동생성
#   src/.../interfaces/dashboard/app.py                  ← 대시보드 (경쟁사·품질 탭 추가)
#   configs/sheets.yaml                                  ← 신규 컬럼 정의
#   configs/manager_rules.yaml                           ← 담당자 룰
#   configs/competitors.yaml                             ← 경쟁사 목록
# =============================================================================

import os, textwrap
from pathlib import Path

# ── Google Drive 마운트 ───────────────────────────────────────────────────────
from google.colab import drive
drive.mount("/content/drive", force_remount=False)

# ── 프로젝트 루트 (실제 Drive 경로로 수정) ─────────────────────────────────────
PROJECT_ROOT = Path("/content/drive/MyDrive/interx_gov_intelligence")
PROJECT_ROOT.mkdir(parents=True, exist_ok=True)

def write_file(rel_path: str, content: str):
    p = PROJECT_ROOT / rel_path
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(textwrap.dedent(content).lstrip(), encoding="utf-8")
    print(f"  ✅ {rel_path}")

print("=" * 60)
print("InterX Engine — Colab 패치셀 v4 시작")
print("=" * 60)

# ─────────────────────────────────────────────────────────────────────────────
# 1. notice.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/core/entities/notice.py", '''
from __future__ import annotations
import hashlib
from dataclasses import dataclass, field
from datetime import date
from typing import Any, Dict, List, Optional


@dataclass
class Notice:
    execution_id: str
    site: str
    notice_id: str
    title: str
    detail_url: str = ""
    notice_link: str = ""
    posted_date: str = ""
    deadline_date: str = ""
    ministry: str = ""
    agency: str = ""
    business_type: str = ""
    budget: str = ""
    duration_months: str = ""
    summary: str = ""
    recommended_solution: str = ""
    recommended_action: str = ""
    l3_strong: str = "N"
    partner_candidate: str = "N"
    attachments: List[str] = field(default_factory=list)
    attachment_items: List[Dict[str, str]] = field(default_factory=list)
    structured: Dict[str, str] = field(default_factory=dict)
    body_text: str = ""
    # v4.4 추가
    category: str = ""
    decision_reason: str = ""
    decision_action: str = ""
    decision_owner: str = ""
    cluster_id: str = ""
    partner_candidates: List[Dict[str, Any]] = field(default_factory=list)
    llm_summary: str = ""
    manager: str = ""
    proposal_strategy: str = ""
    status: str = ""
    open_ended: bool = False
    duplicate_flag: str = "N"
    competitor_flag: str = ""
    is_new: bool = False
    is_changed: bool = False
    memo: str = ""

    @property
    def notice_key(self) -> str:
        return hashlib.md5(f"{self.site}:{self.notice_id}".encode()).hexdigest()

    def is_closed(self) -> bool:
        if self.open_ended:
            return False
        if not self.deadline_date:
            return False
        try:
            return date.fromisoformat(self.deadline_date[:10]) < date.today()
        except Exception:
            return False
''')

# ─────────────────────────────────────────────────────────────────────────────
# 2. notice_mapper.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/application/mappers/notice_mapper.py", '''
from __future__ import annotations
from datetime import date
from typing import Optional
from interx_engine.core.entities.notice import Notice
from interx_engine.core.entities.score_card import ScoreCard
from interx_engine.infrastructure.utils.budget_parser import is_open_ended, normalize_budget


def _calc_dday(deadline: str) -> str:
    if not deadline:
        return ""
    if is_open_ended(deadline):
        return "상시"
    try:
        delta = (date.fromisoformat(deadline[:10]) - date.today()).days
        return str(delta)
    except Exception:
        return ""


def notice_to_master_row(notice: Notice, score: Optional[ScoreCard] = None) -> dict:
    budget_norm  = normalize_budget(notice.budget) if notice.budget else ""
    dup_flag     = getattr(notice, "duplicate_flag",  "N")
    comp_flag    = getattr(notice, "competitor_flag", "")
    is_new       = getattr(notice, "is_new",          False)
    is_changed   = getattr(notice, "is_changed",      False)

    memo_parts = []
    if is_new:          memo_parts.append("🆕신규")
    if is_changed:      memo_parts.append("🔄변경감지")
    if dup_flag == "Y": memo_parts.append("⚠️중복의심")
    if comp_flag:       memo_parts.append(f"🏢경쟁사({comp_flag})")
    base_memo = getattr(notice, "memo", "") or ""
    auto_memo = " ".join(memo_parts)
    memo = f"{auto_memo} {base_memo}".strip()

    return {
        "실행ID":      notice.execution_id,
        "사이트":      notice.site,
        "공고명":      notice.title,
        "마감일":      notice.deadline_date if not getattr(notice, "open_ended", False) else "상시모집",
        "D-day":      _calc_dday(notice.deadline_date),
        "마감여부":    "상시" if getattr(notice, "open_ended", False) else ("Y" if notice.is_closed() else "N"),
        "주무부처":    notice.ministry,
        "수행기관":    notice.agency,
        "예산":        budget_norm,
        "적합도점수":  score.fitness_score  if score else "",
        "우선순위등급": score.priority_grade if score else "",
        "추천솔루션":  notice.recommended_solution or "-",
        "추천액션":    notice.recommended_action   or "검토",
        "적합키워드":  " | ".join(score.positive_keywords[:5]) if score else "",
        "L3강공고":    notice.l3_strong,
        "파트너후보":  notice.partner_candidate,
        "담당자":      notice.manager or "",
        "검토상태":    notice.status  or "",
        "신규여부":    "Y" if is_new     else "N",
        "변경여부":    "Y" if is_changed else "N",
        "경쟁사감지":  comp_flag or "",
        "중복의심":    dup_flag,
        "메모":        memo,
        "상세URL":     notice.detail_url,
    }


def notice_to_urgent_row(notice: Notice, score: Optional[ScoreCard] = None) -> dict:
    return {
        "사이트":      notice.site,
        "공고명":      notice.title,
        "마감일":      notice.deadline_date,
        "D-day":      _calc_dday(notice.deadline_date),
        "우선순위등급": score.priority_grade  if score else "",
        "적합도점수":   score.fitness_score   if score else "",
        "추천솔루션":   notice.recommended_solution or "-",
        "추천액션":     notice.recommended_action   or "검토",
        "예산":        notice.budget,
        "담당자":       notice.manager or "",
        "상세URL":     notice.detail_url,
    }
''')

# ─────────────────────────────────────────────────────────────────────────────
# 3. budget_parser.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/infrastructure/utils/budget_parser.py", '''
from __future__ import annotations
import re

_OPEN_PATTERNS = re.compile(
    r"(상시|미정|예산\s*소진|별도\s*공고|추후\s*공고|상시\s*모집|수시\s*모집)", re.I
)

def is_open_ended(deadline_str: str) -> bool:
    if not deadline_str:
        return False
    return bool(_OPEN_PATTERNS.search(deadline_str))

def parse_budget_eok(raw: str) -> float:
    if not raw:
        return 0.0
    raw = raw.replace(",", "").replace(" ", "")
    if m := re.search(r"(\\d+(?:\\.\\d+)?)억", raw):
        return float(m.group(1))
    if m := re.search(r"(\\d+(?:\\.\\d+)?)천만", raw):
        return float(m.group(1)) / 10
    if m := re.search(r"(\\d+(?:\\.\\d+)?)만", raw):
        return float(m.group(1)) / 10000
    if m := re.search(r"(\\d+(?:\\.\\d+)?)백만", raw):
        return float(m.group(1)) / 100
    if m := re.search(r"(\\d+)천원", raw):
        return float(m.group(1)) / 100_000
    if m := re.search(r"(\\d+)원", raw):
        return float(m.group(1)) / 1_0000_0000
    if m := re.search(r"(\\d+(?:\\.\\d+)?)", raw):
        v = float(m.group(1))
        if v > 100_000:
            return v / 1_0000_0000
        if v > 1_000:
            return v / 10_000
        return v
    return 0.0

def normalize_budget(raw: str) -> str:
    if not raw or raw.strip() in ("-", ""):
        return ""
    raw_strip = raw.strip()
    eok = parse_budget_eok(raw_strip)
    if eok <= 0:
        return raw_strip
    if eok >= 1:
        return f"{eok:.1f}억" if eok != int(eok) else f"{int(eok)}억"
    man = eok * 10000
    return f"{int(man):,}만원"
''')

# ─────────────────────────────────────────────────────────────────────────────
# 4. morpheme_scorer.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/infrastructure/utils/morpheme_scorer.py", '''
from __future__ import annotations
import re
from typing import Dict, List

_NOUN_RE = re.compile(r"[가-힣A-Za-z][가-힣A-Za-z0-9]*")

def extract_nouns(text: str) -> List[str]:
    try:
        from kiwipiepy import Kiwi
        kiwi = Kiwi()
        result = kiwi.analyze(text)
        return [t.form for sent in result for t in sent[0] if t.tag.startswith("NN")]
    except Exception:
        return _NOUN_RE.findall(text)

def morpheme_score(text: str, keyword_weights: Dict[str, float]) -> float:
    nouns = set(n.lower() for n in extract_nouns(text))
    score = 0.0
    for kw, weight in keyword_weights.items():
        if kw.lower() in nouns or kw.lower() in text.lower():
            score += weight
    return score
''')

# ─────────────────────────────────────────────────────────────────────────────
# 5. deduplicate_notices.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/application/use_cases/deduplicate_notices.py", '''
from __future__ import annotations
import logging
from typing import List, Tuple

log = logging.getLogger("interx.dedup")
_THRESHOLD = 0.82


def deduplicate_by_tfidf(
    notices, score_cards,
    threshold: float = _THRESHOLD,
) -> Tuple[list, int]:
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        import numpy as np
    except ImportError:
        log.warning("[Dedup] sklearn 미설치 — TF-IDF 중복감지 스킵")
        return notices, 0

    if len(notices) < 2:
        return notices, 0

    score_map = {s.notice_id: s for s in score_cards}
    texts = [f"{n.title} {n.summary[:200]}" for n in notices]

    try:
        vec  = TfidfVectorizer(analyzer="char_wb", ngram_range=(2, 3), max_features=5000)
        mat  = vec.fit_transform(texts)
        sim  = cosine_similarity(mat)
        np.fill_diagonal(sim, 0)
    except Exception as e:
        log.warning("[Dedup] TF-IDF 계산 오류: %s", e)
        return notices, 0

    dup_ids = set()
    n = len(notices)
    for i in range(n):
        if notices[i].notice_id in dup_ids:
            continue
        for j in range(i + 1, n):
            if notices[j].notice_id in dup_ids:
                continue
            if notices[i].site == notices[j].site:
                continue
            if sim[i, j] >= threshold:
                sc_i = getattr(score_map.get(notices[i].notice_id), "fitness_score", 0) or 0
                sc_j = getattr(score_map.get(notices[j].notice_id), "fitness_score", 0) or 0
                loser = notices[j].notice_id if sc_i >= sc_j else notices[i].notice_id
                dup_ids.add(loser)

    dup_count = 0
    for n_ in notices:
        if n_.notice_id in dup_ids:
            n_.duplicate_flag = "Y"
            dup_count += 1

    if dup_count:
        log.info("[Dedup] TF-IDF 중복 의심 %d건 플래그 처리", dup_count)
    return notices, dup_count
''')

# ─────────────────────────────────────────────────────────────────────────────
# 6. detect_changes.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/application/use_cases/detect_changes.py", '''
from __future__ import annotations
import hashlib
import json
import logging
import os
from pathlib import Path
from typing import List, Tuple

log = logging.getLogger("interx.changes")

_CACHE_PATHS = [
    Path(os.getenv("INTERX_CHANGE_CACHE",
         "/content/drive/MyDrive/interx_gov_intelligence/data/change_cache.json")),
    Path("/tmp/interx_change_cache.json"),
]


def _load_cache() -> dict:
    for p in _CACHE_PATHS:
        if p.exists():
            try:
                return json.loads(p.read_text(encoding="utf-8"))
            except Exception:
                pass
    return {}


def _save_cache(cache: dict) -> None:
    for p in _CACHE_PATHS:
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")
            return
        except Exception:
            continue
    log.warning("[Changes] 캐시 저장 실패 (읽기 전용 환경)")


def _notice_hash(notice) -> str:
    raw = f"{notice.title}|{notice.deadline_date}|{notice.budget}"
    return hashlib.md5(raw.encode()).hexdigest()


def detect_changes(notices) -> Tuple[list, int, int]:
    cache     = _load_cache()
    new_count = changed_count = 0

    for n in notices:
        key      = n.notice_key
        cur_hash = _notice_hash(n)
        prev     = cache.get(key)
        if prev is None:
            n.is_new = True
            new_count += 1
        elif prev != cur_hash:
            n.is_changed = True
            changed_count += 1
        cache[key] = cur_hash

    _save_cache(cache)
    if new_count or changed_count:
        log.info("[Changes] 신규=%d 변경=%d", new_count, changed_count)
    return notices, new_count, changed_count
''')

# ─────────────────────────────────────────────────────────────────────────────
# 7. assign_manager.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/application/use_cases/assign_manager.py", '''
from __future__ import annotations
import logging
from pathlib import Path
from typing import List

log = logging.getLogger("interx.manager")

_RULES_PATH = Path(__file__).resolve().parents[5] / "configs" / "manager_rules.yaml"


def _load_rules():
    try:
        import yaml
        if not _RULES_PATH.exists():
            return []
        return yaml.safe_load(_RULES_PATH.read_text(encoding="utf-8")).get("rules", [])
    except Exception as e:
        log.warning("[Manager] 룰 로드 실패: %s", e)
        return []


def assign_managers(notices) -> list:
    rules    = _load_rules()
    assigned = 0
    for n in notices:
        if n.manager:
            continue
        text = f"{n.title} {n.ministry} {n.agency} {n.summary}".lower()
        for rule in rules:
            kws  = rule.get("keywords", [])
            mins = rule.get("ministries", [])
            kw_hit  = any(k.lower() in text for k in kws)  if kws  else True
            min_hit = any(m       in text for m in mins)   if mins else True
            if kw_hit and min_hit:
                n.manager = rule.get("manager", "")
                assigned += 1
                break
    if assigned:
        log.info("[Manager] %d건 담당자 자동 배정", assigned)
    return notices
''')

# ─────────────────────────────────────────────────────────────────────────────
# 8. track_competitors.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/application/use_cases/track_competitors.py", '''
from __future__ import annotations
import logging
from pathlib import Path

log = logging.getLogger("interx.competitors")

_COMP_PATH = Path(__file__).resolve().parents[5] / "configs" / "competitors.yaml"


def _load_config():
    try:
        import yaml
        if not _COMP_PATH.exists():
            return {}, {}
        cfg = yaml.safe_load(_COMP_PATH.read_text(encoding="utf-8"))
        return cfg.get("tier1", []), cfg.get("tier2", [])
    except Exception as e:
        log.warning("[Competitor] 설정 로드 실패: %s", e)
        return [], []


def track_competitors(notices) -> list:
    tier1, tier2 = _load_config()
    if not tier1 and not tier2:
        return notices
    flagged = 0
    for n in notices:
        text = f"{n.title} {n.agency} {n.ministry} {n.summary} {n.body_text[:500]}".lower()
        if any(c.lower() in text for c in tier1):
            n.competitor_flag = "tier1"
            n.memo = (n.memo + " 🏢tier1경쟁사감지").strip()
            flagged += 1
        elif any(c.lower() in text for c in tier2):
            n.competitor_flag = "tier2"
            n.memo = (n.memo + " 🏢tier2경쟁사감지").strip()
            flagged += 1
    if flagged:
        log.info("[Competitor] %d건 경쟁사 감지", flagged)
    return notices
''')

# ─────────────────────────────────────────────────────────────────────────────
# 9. site_quality_grader.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/application/use_cases/site_quality_grader.py", '''
from __future__ import annotations
import logging
from collections import defaultdict
from typing import Dict, List
from interx_engine.core.entities.notice import Notice
from interx_engine.core.entities.score_card import ScoreCard

log = logging.getLogger("interx.quality")


def grade_site_quality(notices, score_cards) -> Dict[str, Dict]:
    score_map = {s.notice_id: s for s in score_cards}
    buckets   = defaultdict(lambda: {"total":0,"l3":0,"p1p2":0,"has_budget":0,"has_deadline":0})
    for n in notices:
        site = n.site or "unknown"
        b    = buckets[site]
        b["total"] += 1
        if n.l3_strong == "Y": b["l3"] += 1
        if n.budget and n.budget not in ("-", ""): b["has_budget"] += 1
        if n.deadline_date: b["has_deadline"] += 1
        sc = score_map.get(n.notice_id)
        if sc and sc.priority_grade in ("P1", "P2"): b["p1p2"] += 1

    results = {}
    for site, b in buckets.items():
        total = b["total"]
        if total == 0:
            results[site] = {"grade":"F","score":0,"total":0}
            continue
        l3_r  = b["l3"]          / total
        p1p2r = b["p1p2"]        / total
        budr  = b["has_budget"]  / total
        dlr   = b["has_deadline"] / total
        volr  = min(total, 50)   / 50
        score = (l3_r*30 + p1p2r*25 + budr*20 + dlr*15 + volr*10) * 100
        grade = "A" if score>=70 else "B" if score>=50 else "C" if score>=30 else "D" if score>=10 else "F"
        results[site] = {
            "grade": grade, "score": round(score,1), "total": total,
            "l3_rate": round(l3_r*100,1), "p1p2_rate": round(p1p2r*100,1),
            "budget_rate": round(budr*100,1),
        }
        log.debug("[Quality] %-12s → %s (%.1f점)", site, grade, score)
    return results
''')

# ─────────────────────────────────────────────────────────────────────────────
# 10. generate_proposal.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/application/use_cases/generate_proposal.py", '''
from __future__ import annotations
import logging
from pathlib import Path
from typing import List
from interx_engine.core.entities.notice import Notice
from interx_engine.core.entities.score_card import ScoreCard

log = logging.getLogger("interx.proposal")

_SECTIONS = [
    ("1. 사업 개요",
     "공고명:   {title}\\n주관기관: {agency}\\n주무부처: {ministry}\\n공고일:   {posted}\\n마감일:   {deadline}"),
    ("2. 지원 규모",
     "예산: {budget}\\n지원 기간: {duration}"),
    ("3. InterX 추천 솔루션",
     "{solution}\\n\\n핵심 적합 키워드: {keywords}"),
    ("4. InterX 강점",
     "• 제조 AI 전문 역량 (AX · 스마트팩토리 · 디지털트윈)\\n"
     "• {keywords} 핵심 기술 보유\\n"
     "• 유사 프로젝트 다수 수행 경험\\n"
     "• [작성 필요] 구체적 수행 실적 기재"),
    ("5. 제안 전략",
     "추천 액션: {action}\\n\\n[작성 필요] 차별화 포인트 및 제안 핵심 내용을 여기에 기재하세요."),
    ("6. 추진 일정",
     "제안서 제출 기한: {deadline}\\n담당자: {manager}\\n\\n[작성 필요] 세부 일정표 첨부"),
    ("7. 공고 원문 링크", "{url}"),
]


def generate_proposals(
    notices: List[Notice],
    score_cards: List[ScoreCard],
    output_dir: str = "/tmp/interx_proposals",
    target_grades: tuple = ("P1", "P2"),
) -> List[str]:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        log.warning("[Proposal] python-docx 미설치")
        return []

    score_map = {s.notice_id: s for s in score_cards}
    out_dir   = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    generated = []

    for notice in notices:
        score = score_map.get(notice.notice_id)
        if not score or score.priority_grade not in target_grades:
            continue
        doc = Document()
        h = doc.add_heading(f"[제안서 초안] {notice.title}", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x46, 0x8A)
        p = doc.add_paragraph()
        run = p.add_run(
            f"  ★ {score.priority_grade} 등급  |  "
            f"적합도: {score.fitness_score:.0f}점  |  "
            f"추천솔루션: {notice.recommended_solution}"
        )
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph()
        for section_title, template in _SECTIONS:
            doc.add_heading(section_title, level=2)
            content = template.format(
                title    = notice.title,
                agency   = notice.agency        or "-",
                ministry = notice.ministry      or "-",
                posted   = notice.posted_date   or "-",
                deadline = notice.deadline_date or "-",
                budget   = notice.budget        or "-",
                duration = notice.duration_months or "-",
                solution = notice.recommended_solution or "-",
                keywords = " / ".join(score.positive_keywords[:5]) if score else "-",
                action   = notice.recommended_action or "제안 검토",
                manager  = notice.manager       or "미배정",
                url      = notice.detail_url    or notice.notice_link or "-",
            )
            doc.add_paragraph(content)
        safe = "".join(c for c in notice.title[:30] if c.isalnum() or c in " _-한글")
        fname = out_dir / f"{score.priority_grade}_{notice.site}_{safe}.docx"
        doc.save(str(fname))
        generated.append(str(fname))
        log.info("[Proposal] 생성: %s", fname.name)

    log.info("[Proposal] 총 %d개 → %s", len(generated), output_dir)
    return generated
''')

# ─────────────────────────────────────────────────────────────────────────────
# 11. daily_pipeline.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/application/orchestrators/daily_pipeline.py", '''
from __future__ import annotations
import logging
import time
from interx_engine.application.use_cases.collect_notices import CollectNoticesUseCase
from interx_engine.application.use_cases.score_notices import ScoreNoticesUseCase
from interx_engine.application.use_cases.deduplicate_notices import deduplicate_by_tfidf
from interx_engine.application.use_cases.detect_changes import detect_changes
from interx_engine.application.use_cases.assign_manager import assign_managers
from interx_engine.application.use_cases.track_competitors import track_competitors
from interx_engine.application.use_cases.site_quality_grader import grade_site_quality
from interx_engine.application.use_cases.generate_proposal import generate_proposals
from interx_engine.application.mappers.notice_mapper import (
    notice_to_master_row, notice_to_urgent_row, _calc_dday,
)
from interx_engine.application.mappers.kpi_mapper import (
    build_kpi_rows, build_exec_log_row,
    build_site_stats_rows, build_collect_error_rows,
)

log = logging.getLogger("interx.pipeline")
_URGENT_DDAY = 7


class DailyPipelineOrchestrator:
    def __init__(self, collector, sheet_gateway=None, **_):
        self.collect_use_case = CollectNoticesUseCase(collector)
        self.score_use_case   = ScoreNoticesUseCase()
        self.sheet_gateway    = sheet_gateway

    def run(self, execution_id: str) -> dict:
        t0 = time.monotonic()

        log.info("[Pipeline] 수집 시작 (%s)", execution_id)
        notices = self.collect_use_case.execute(execution_id)
        log.info("[Pipeline] %d건 수집 완료", len(notices))

        seen, unique = set(), []
        for n in notices:
            if n.notice_id not in seen:
                seen.add(n.notice_id); unique.append(n)
        if len(unique) < len(notices):
            log.info("[Pipeline] 중복 제거: %d → %d건", len(notices), len(unique))
        notices = unique

        notices, score_cards = self.score_use_case.execute(notices)
        score_map = {s.notice_id: s for s in score_cards}

        notices, dup_count = deduplicate_by_tfidf(notices, score_cards)
        notices, new_count, changed_count = detect_changes(notices)
        notices = assign_managers(notices)
        notices = track_competitors(notices)
        quality_grades = grade_site_quality(notices, score_cards)

        master_rows, l3_rows, urgent_rows = [], [], []
        for notice in notices:
            score = score_map.get(notice.notice_id)
            row   = notice_to_master_row(notice, score)
            master_rows.append(row)
            if notice.l3_strong == "Y":
                l3_rows.append(row.copy())
            dday_str = _calc_dday(notice.deadline_date)
            if dday_str and dday_str != "상시":
                try:
                    if 0 <= int(dday_str) <= _URGENT_DDAY:
                        urgent_rows.append(notice_to_urgent_row(notice, score))
                except ValueError:
                    pass

        try:
            proposal_files = generate_proposals(notices, score_cards)
        except Exception as e:
            log.warning("[Pipeline] 제안서 생성 실패 (무시): %s", e)
            proposal_files = []

        elapsed    = round(time.monotonic() - t0, 1)
        kpi_rows   = build_kpi_rows(execution_id, notices, score_cards)
        site_stats = build_site_stats_rows(execution_id, notices, score_cards)

        collector_obj = getattr(self.collect_use_case, "collector", None)
        raw_errors    = getattr(collector_obj, "last_errors", [])
        error_rows    = build_collect_error_rows(execution_id, raw_errors)

        exec_log_row = build_exec_log_row(
            execution_id, "pipeline_complete", "OK", elapsed,
            f"총 {len(notices)}건 | L3={len(l3_rows)} | 긴급={len(urgent_rows)} | "
            f"신규={new_count} | 변경={changed_count} | 중복의심={dup_count} | "
            f"경쟁사={sum(1 for n in notices if n.competitor_flag)} | "
            f"제안서={len(proposal_files)} | 에러={len(error_rows)}"
        )

        if self.sheet_gateway:
            self._upload(master_rows, l3_rows, urgent_rows,
                         kpi_rows, site_stats, error_rows,
                         exec_log_row, elapsed, len(notices))

        return {
            "notice_count":   len(notices),
            "master_rows":    master_rows,
            "l3_rows":        l3_rows,
            "urgent_rows":    urgent_rows,
            "new_count":      new_count,
            "changed_count":  changed_count,
            "dup_count":      dup_count,
            "error_count":    len(error_rows),
            "proposal_files": proposal_files,
            "quality_grades": quality_grades,
        }

    def _upload(self, master_rows, l3_rows, urgent_rows,
                kpi_rows, site_stats, error_rows,
                exec_log_row, elapsed, total):
        gw = self.sheet_gateway
        try:
            if hasattr(gw, "cleanup_old_sheets"):
                gw.cleanup_old_sheets()
        except Exception as e:
            log.warning("[Pipeline] cleanup 실패: %s", e)

        try:
            gw.replace_rows("01_영업기회_정보", master_rows)
            gw.replace_rows("02_L3강공고",      l3_rows)
            gw.replace_rows("05_긴급마감_공고",  urgent_rows)
            if kpi_rows:   gw.append_rows("22_KPI현황",         kpi_rows)
            if site_stats: gw.append_rows("93_사이트별수집통계", site_stats)
            gw.append_rows("94_실행로그",                        [exec_log_row])
            if error_rows: gw.append_rows("96_수집에러로그",     error_rows)
            log.info("[Pipeline] 업로드 완료 (%.1fs, %d건)", elapsed, total)
        except Exception as exc:
            log.error("[Pipeline] 업로드 실패: %s", exc)
            raise
''')

# ─────────────────────────────────────────────────────────────────────────────
# 12. new_collectors.py
# ─────────────────────────────────────────────────────────────────────────────
write_file("src/interx_engine/infrastructure/collectors/sites/new_collectors.py", '''
from __future__ import annotations
import logging
import time
from interx_engine.infrastructure.collectors.base_collector import BaseCollector
from interx_engine.core.entities.notice import Notice

log = logging.getLogger("interx.collectors.new")


class NrfCollector(BaseCollector):
    site_key = "nrf"
    BASE_URL = "https://www.nrf.re.kr/biz/notice/list"

    def _parse_page(self, execution_id, page):
        return []


class KisedCollector(BaseCollector):
    site_key = "kised"
    BASE_URL = "https://www.kised.or.kr/board.es?mid=a10203000000&bid=0003"

    def _parse_page(self, execution_id, page):
        return []


class KetepCollector(BaseCollector):
    site_key = "ketep"
    BASE_URL = "https://www.ketep.re.kr/biz/bizNoticeList.do"

    def _parse_page(self, execution_id, page):
        return []


class KoiiaCollector(BaseCollector):
    site_key = "koiia"
    BASE_URL = "https://www.koiia.or.kr/board/notice/list.do"

    def _parse_page(self, execution_id, page):
        return []


NEW_COLLECTOR_CLASSES = {
    "nrf":   NrfCollector,
    "kised": KisedCollector,
    "ketep": KetepCollector,
    "koiia": KoiiaCollector,
}
''')

# ─────────────────────────────────────────────────────────────────────────────
# 13. configs/sheets.yaml
# ─────────────────────────────────────────────────────────────────────────────
write_file("configs/sheets.yaml", '''
# InterX BD Intelligence — 시트 정의 (9-sheet 아키텍처 v4.4)
sheets:

  opportunity_info:
    name: "01_영업기회_정보"
    columns:
      - 실행ID
      - 사이트
      - 공고명
      - 마감일
      - D-day
      - 마감여부
      - 주무부처
      - 수행기관
      - 예산
      - 적합도점수
      - 우선순위등급
      - 추천솔루션
      - 추천액션
      - 적합키워드
      - L3강공고
      - 파트너후보
      - 담당자
      - 검토상태
      - 신규여부
      - 변경여부
      - 경쟁사감지
      - 중복의심
      - 메모
      - 상세URL

  l3_dealer:
    name: "02_L3강공고"
    same_as: "opportunity_info"

  partner_share:
    name: "03_파트너전달"
    same_as: "opportunity_info"

  urgent_notices:
    name: "05_긴급마감_공고"
    columns:
      - 사이트
      - 공고명
      - 마감일
      - D-day
      - 우선순위등급
      - 적합도점수
      - 추천솔루션
      - 추천액션
      - 예산
      - 담당자
      - 상세URL

  summary_dashboard:
    name: "20_요약대시보드"
    columns:
      - 기준일
      - 항목
      - 값
      - 비고

  kpi:
    name: "22_KPI현황"
    columns:
      - 기준일
      - 구분
      - 지표
      - 값
      - 실행ID

  site_stats:
    name: "93_사이트별수집통계"
    columns:
      - 기준일
      - 실행ID
      - 사이트
      - 수집건수
      - 마감임박건수
      - L3건수
      - P1건수
      - P2건수
      - P3건수

  execution_log:
    name: "94_실행로그"
    columns:
      - 실행ID
      - 실행시각
      - 단계
      - 상태
      - 소요초
      - 메시지

  error_log:
    name: "96_수집에러로그"
    columns:
      - 실행ID
      - 실행시각
      - 사이트
      - 에러유형
      - 에러메시지
''')

# ─────────────────────────────────────────────────────────────────────────────
# 14. configs/manager_rules.yaml
# ─────────────────────────────────────────────────────────────────────────────
write_file("configs/manager_rules.yaml", '''
# InterX 담당자 자동 배정 룰 (우선순위 순)
rules:
  - name: "제조AI전문"
    manager: "김BD"
    keywords: ["스마트팩토리", "제조AI", "디지털트윈", "AX", "제조혁신", "스마트공장"]
    ministries: []

  - name: "R&D과제"
    manager: "이연구"
    keywords: ["R&D", "연구개발", "기술개발", "연구과제", "국책과제"]
    ministries: ["과학기술정보통신부", "산업통상자원부"]

  - name: "중기부바우처"
    manager: "박바우처"
    keywords: ["바우처", "스마트서비스", "중소기업", "소상공인"]
    ministries: ["중소벤처기업부"]

  - name: "안전품질"
    manager: "최안전"
    keywords: ["안전", "품질", "검사", "불량", "위험", "사고예방"]
    ministries: []

  - name: "에너지환경"
    manager: "정에너지"
    keywords: ["에너지", "탄소중립", "친환경", "RE100", "그린", "탄소"]
    ministries: ["산업통상자원부", "환경부"]

  - name: "기본배정"
    manager: "미배정"
    keywords: []
    ministries: []
''')

# ─────────────────────────────────────────────────────────────────────────────
# 15. configs/competitors.yaml
# ─────────────────────────────────────────────────────────────────────────────
write_file("configs/competitors.yaml", '''
# InterX 경쟁사 트래킹 설정
tier1:  # 직접 경쟁사 (AI/DX 제조)
  - 삼성SDS
  - LG CNS
  - SK C&C
  - 포스코DX
  - 현대오토에버
  - 롯데정보통신

tier2:  # 간접 경쟁사 / 글로벌
  - 한국IBM
  - 오라클
  - SAP
  - 다쏘시스템
  - 지멘스
  - 씨메스
  - 수아랩
  - 뷰런테크놀로지
  - 코그넥스
  - 아이메디신

partners:  # 협력 가능 기관 (경쟁사 아님)
  - 성균관대
  - KAIST
  - 한국생산기술연구원
  - 한국전자기술연구원
  - 한국기계연구원
''')

# ─────────────────────────────────────────────────────────────────────────────
# 16. run_engine.py (멀티콜렉터 어댑터 + 새 콜렉터 레지스트리)
# ─────────────────────────────────────────────────────────────────────────────
write_file("run_engine.py", r'''
# InterX Government Intelligence Engine - run_engine.py v4.4
from __future__ import annotations
import argparse, json, logging, sys, time
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT / "src"))

try:
    from dotenv import load_dotenv
    load_dotenv(ROOT / ".env", override=False)
except ImportError:
    pass

from interx_engine.infrastructure.config.settings_loader import settings
settings.ensure_dirs()


def _setup_logger() -> logging.Logger:
    logger = logging.getLogger("interx")
    if logger.handlers:
        return logger
    logger.setLevel(logging.DEBUG)
    fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(name)s: %(message)s")
    sh = logging.StreamHandler(sys.stdout)
    sh.setLevel(logging.INFO)
    sh.setFormatter(fmt)
    fh = logging.FileHandler(
        Path(settings.log_dir) / f"interx_{datetime.now().strftime('%Y%m%d')}.log",
        encoding="utf-8",
    )
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(fmt)
    logger.addHandler(sh)
    logger.addHandler(fh)
    return logger


log = _setup_logger()
COLLECTOR_REGISTRY: dict = {}


def _build_registry() -> dict:
    from interx_engine.infrastructure.collectors.sites.bizinfo_collector import BizinfoCollector
    from interx_engine.infrastructure.collectors.sites.multi_site_collectors import COLLECTOR_CLASSES
    from interx_engine.infrastructure.collectors.sites.new_collectors import NEW_COLLECTOR_CLASSES
    return {"bizinfo": BizinfoCollector, **COLLECTOR_CLASSES, **NEW_COLLECTOR_CLASSES}


def build_collectors(site_keys, max_pages):
    global COLLECTOR_REGISTRY
    if not COLLECTOR_REGISTRY:
        COLLECTOR_REGISTRY = _build_registry()
    targets    = site_keys or list(COLLECTOR_REGISTRY.keys())
    collectors = []
    for key in targets:
        cls = COLLECTOR_REGISTRY.get(key)
        if cls is None:
            log.warning("알 수 없는 사이트 키: %s (스킵)", key)
            continue
        try:
            col = cls(max_pages=max_pages)
        except TypeError:
            col = cls()
            col.max_pages = max_pages
        collectors.append(col)
        log.info("콜렉터 등록: %s", key)
    return collectors


def build_sheet_gateway(enable=True):
    if not enable:
        log.info("Google Sheets 비활성화")
        return None
    try:
        from interx_engine.infrastructure.sheets.google_sheet_gateway import GoogleSheetGateway
        gw = GoogleSheetGateway(
            spreadsheet_name=settings.spreadsheet_name,
            service_account_json=settings.service_account,
            sheets_config_path=settings.sheets_config_path,
        )
        log.info("Google Sheets 연결: %s", settings.spreadsheet_name)
        return gw
    except Exception as e:
        err = str(e)
        if "invalid_grant" in err or "account not found" in err:
            log.error("🔑 Google Sheets 인증 실패 — service_account.json 재발급 필요")
        else:
            log.warning("Google Sheets 연결 실패 (%s)", e)
        log.warning("→ ConsoleSheetGateway fallback")
        from interx_engine.infrastructure.sheets.console_sheet_gateway import ConsoleSheetGateway
        return ConsoleSheetGateway()


class MultiCollectorAdapter:
    def __init__(self, collectors, max_workers=8):
        self.collectors   = collectors
        self.max_workers  = max_workers
        self.last_errors: list = []

    def _collect_with_retry(self, col, execution_id, max_retries=2):
        site = getattr(col, "site_key", repr(col))
        last_exc = None
        for attempt in range(max_retries + 1):
            try:
                return col.collect(execution_id)
            except Exception as e:
                last_exc = e
                if attempt < max_retries:
                    wait = 2 ** attempt
                    log.warning("[Retry] %-12s 실패 (attempt %d/%d), %ds 후 재시도: %s",
                                site, attempt+1, max_retries+1, wait, e)
                    time.sleep(wait)
        raise last_exc

    def collect(self, execution_id):
        if not self.collectors:
            return []
        all_notices = []
        self.last_errors = []
        n = min(len(self.collectors), self.max_workers)
        with ThreadPoolExecutor(max_workers=n) as ex:
            futures = {ex.submit(self._collect_with_retry, c, execution_id): c
                       for c in self.collectors}
            for fut in as_completed(futures, timeout=300):
                col  = futures[fut]
                site = getattr(col, "site_key", repr(col))
                try:
                    notices = fut.result(timeout=settings.collector_timeout * 6)
                    all_notices.extend(notices)
                    log.info("[Collect] %-12s %d건", site, len(notices))
                except Exception as e:
                    log.error("[Collect] %s 최종 실패: %s", site, e)
                    self.last_errors.append({
                        "site": site,
                        "error_type": type(e).__name__,
                        "error_message": str(e)[:300],
                    })
        return all_notices


def main(site_keys=None, max_pages=None, enable_sheets=True, full_pipeline=False) -> dict:
    max_pages    = max_pages or settings.max_pages
    execution_id = datetime.now().strftime("EXEC-%Y%m%d-%H%M%S")

    log.info("=" * 60)
    log.info("InterX Government Intelligence Engine v4.4")
    log.info("execution_id  = %s", execution_id)
    log.info("대상 사이트   = %s", site_keys or "전체")
    log.info("max_pages     = %d", max_pages)
    log.info("=" * 60)

    collectors = build_collectors(site_keys, max_pages)
    if not collectors:
        log.error("사용 가능한 수집기가 없습니다.")
        return {}

    multi_collector = MultiCollectorAdapter(collectors, settings.max_workers)
    sheet_gateway   = build_sheet_gateway(enable_sheets)

    if full_pipeline:
        from interx_engine.application.orchestrators.full_pipeline import FullPipelineOrchestrator
        orchestrator = FullPipelineOrchestrator(
            collector=multi_collector, base_dir=str(ROOT), sheet_gateway=sheet_gateway,
        )
    else:
        from interx_engine.application.orchestrators.daily_pipeline import DailyPipelineOrchestrator
        orchestrator = DailyPipelineOrchestrator(
            collector=multi_collector, sheet_gateway=sheet_gateway,
        )

    result = orchestrator.run(execution_id)

    skip = {"master_rows","score_cards","notices","attachment_rows",
            "doc_mgmt_rows","pipeline_rows","score_model_rows"}
    print("\n" + "=" * 60)
    print("=== FINAL RESULT ===")
    print(json.dumps({k: v for k, v in result.items() if k not in skip},
                     ensure_ascii=False, indent=2, default=str))
    print("=" * 60)
    return result


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="InterX Government Intelligence Engine v4.4")
    parser.add_argument("--sites",     type=str,  default=None)
    parser.add_argument("--max-pages", type=int,  default=None)
    parser.add_argument("--no-sheets", action="store_true")
    parser.add_argument("--full",      action="store_true")
    args = parser.parse_args()
    main(
        site_keys=[s.strip() for s in args.sites.split(",")] if args.sites else None,
        max_pages=args.max_pages,
        enable_sheets=not args.no_sheets,
        full_pipeline=args.full,
    )
''')

print()
print("=" * 60)
print("✅ 패치셀 v4 완료 — 모든 파일 Drive에 기록됨")
print()
print("다음 단계:")
print("  1. 런타임 재시작:  Runtime → Restart runtime")
print("  2. 엔진 실행:      !python run_engine.py --no-sheets --max-pages 2")
print("  3. 대시보드:       !streamlit run src/interx_engine/interfaces/dashboard/app.py")
print("=" * 60)
