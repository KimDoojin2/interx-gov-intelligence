"""
테스트용 샘플 PDF/DOCX 파일 생성.
실행: venv/Scripts/python tests/fixtures/generate_samples.py
"""
import io
import struct
import zlib
from pathlib import Path

HERE = Path(__file__).parent


def generate_sample_pdf() -> bytes:
    """reportlab로 한글 PDF 생성."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    pdfmetrics.registerFont(TTFont("Malgun", "C:/Windows/Fonts/malgun.ttf"))
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("Malgun", 12)

    lines = [
        "2026년 스마트공장 구축 지원사업 공고",
        "",
        "1. 사업목적",
        "  중소기업 제조 현장에 스마트공장 기술을 보급하여",
        "  생산성 향상 및 품질 개선을 지원합니다.",
        "",
        "2. 지원대상",
        "  제조업 중소기업 (매출 1000억원 이하)",
        "",
        "3. 지원내용",
        "  MES, ERP, AI 품질검사, 디지털트윈 구축비 지원",
        "  설비예지보전(PdM) 시스템 도입 지원",
        "",
        "4. 지원금액: 최대 3억원 (정부 50% + 자부담 50%)",
        "",
        "5. 신청기간: 2026.06.01 ~ 2026.06.30",
        "",
        "6. 문의: 스마트제조혁신추진단 02-1234-5678",
    ]
    y = 750
    for line in lines:
        c.drawString(72, y, line)
        y -= 18
    c.save()
    return buf.getvalue()


def generate_sample_docx() -> bytes:
    """python-docx로 한글 DOCX 생성."""
    from docx import Document

    doc = Document()
    doc.add_heading("2026년 제조 AI 도입 지원사업 공고", level=1)
    doc.add_paragraph("")
    doc.add_heading("1. 사업목적", level=2)
    doc.add_paragraph(
        "중소 제조기업의 AI 기반 품질검사 및 공정최적화 시스템 구축을 지원하여 "
        "제조 경쟁력을 강화하고 스마트 제조 전환을 촉진합니다."
    )
    doc.add_heading("2. 지원대상", level=2)
    doc.add_paragraph("제조업 중소기업 (업력 3년 이상, 매출 500억원 이하)")
    doc.add_heading("3. 지원내용", level=2)
    doc.add_paragraph("- AI 품질검사 시스템 구축")
    doc.add_paragraph("- 설비예지보전(PdM) 도입")
    doc.add_paragraph("- 디지털트윈 플랫폼 구축")
    doc.add_paragraph("- MES/ERP 고도화")
    doc.add_heading("4. 지원금액", level=2)
    doc.add_paragraph("최대 5억원 (정부 70% + 자부담 30%)")
    doc.add_heading("5. 신청기간", level=2)
    doc.add_paragraph("2026.07.01 ~ 2026.07.31")

    # 테이블 추가
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    headers = ["구분", "지원항목", "지원한도"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows_data = [
        ["기본형", "MES + ERP", "2억원"],
        ["고도화형", "AI + 디지털트윈", "5억원"],
    ]
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_sample_hwp_binary() -> bytes:
    """HWP 바이너리 레코드 (Section 데이터) 생성 — olefile OLE 컨테이너 없이."""
    def make_record(tag_id: int, data: bytes) -> bytes:
        size = len(data)
        if size >= 0xFFF:
            header = (tag_id & 0x3FF) | (0xFFF << 20)
            return struct.pack("<I", header) + struct.pack("<I", size) + data
        else:
            header = (tag_id & 0x3FF) | ((size & 0xFFF) << 20)
            return struct.pack("<I", header) + data

    paragraphs = [
        "2026년 스마트공장 구축 지원사업 공고문입니다.",
        "사업목적: 중소기업 제조 현장에 스마트팩토리 기술을 보급합니다.",
        "지원대상: 제조업 중소기업 매출 1000억원 이하",
        "지원내용: MES, ERP, AI 품질검사, 디지털트윈 구축비 지원",
        "지원금액: 최대 3억원 정부 50퍼센트 자부담 50퍼센트",
    ]

    records = b""
    for para in paragraphs:
        text_bytes = para.encode("utf-16-le")
        records += make_record(67, text_bytes)  # HWPTAG_PARA_TEXT = 67
        records += make_record(10, b"\x00" * 8)  # dummy non-text record

    return records


if __name__ == "__main__":
    pdf = generate_sample_pdf()
    (HERE / "sample_notice.pdf").write_bytes(pdf)
    print(f"PDF: {len(pdf)} bytes -> tests/fixtures/sample_notice.pdf")

    docx = generate_sample_docx()
    (HERE / "sample_notice.docx").write_bytes(docx)
    print(f"DOCX: {len(docx)} bytes -> tests/fixtures/sample_notice.docx")

    hwp = generate_sample_hwp_binary()
    (HERE / "sample_hwp_section.bin").write_bytes(hwp)
    print(f"HWP binary: {len(hwp)} bytes -> tests/fixtures/sample_hwp_section.bin")
